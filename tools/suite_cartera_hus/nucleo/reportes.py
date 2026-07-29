"""
nucleo.reportes — Estándar único de evidencia para TODOS los procesos de la Suite.

Contrato de salida (el mismo que ya usan COOSALUD y Dispensario Médico):

    SALIDAS/<ENTIDAD>/<AAAA-MM-DD_HHMM>_<PROCESO>/
        ├── CONTROL.csv        ← una fila por factura/ítem con su estado
        ├── EVIDENCIAS.docx    ← un pantallazo por página (si hay imágenes)
        ├── REVISAR.csv        ← lo excluido, con su motivo (trazabilidad)
        └── RESUMEN.txt        ← cifras del cierre, legible para el informe

Columnas fijas de CONTROL.csv (no cambiar: los tableros las leen así):
    fecha_hora | entidad | proceso | factura | estado | detalle | valor | evidencia

Uso:
    rep = ReporteEstandar("COOSALUD", "respuesta_glosas", base="SALIDAS")
    rep.registrar("HUS123", "CERRADA", "24 glosas respondidas", 1500000, "HUS123.png")
    rep.excluir("HUS999", "Factura ya objetada en DGH")
    rep.compilar_evidencias_word("carpeta/con/pantallazos")   # opcional
    rep.cerrar()   # escribe CONTROL.csv, REVISAR.csv y RESUMEN.txt
"""

import csv
import os
import re
from datetime import datetime

COLUMNAS_CONTROL = [
    "fecha_hora",
    "entidad",
    "proceso",
    "factura",
    "estado",
    "detalle",
    "valor",
    "evidencia",
]
COLUMNAS_REVISAR = ["fecha_hora", "entidad", "proceso", "factura", "motivo", "detalle"]


def _slug(texto):
    limpio = re.sub(r"[^A-Za-z0-9_-]+", "_", str(texto).strip())
    return limpio.strip("_")[:60] or "SIN_NOMBRE"


class ReporteEstandar(object):
    def __init__(self, entidad, proceso, base="SALIDAS", log=print):
        self.entidad = str(entidad)
        self.proceso = str(proceso)
        self.log = log
        marca = datetime.now().strftime("%Y-%m-%d_%H%M")
        self.carpeta = os.path.join(base, _slug(entidad), "%s_%s" % (marca, _slug(proceso)))
        os.makedirs(self.carpeta, exist_ok=True)
        self._filas = []
        self._revisar = []
        self._inicio = datetime.now()

    # ------------------------------------------------------------ registro

    def registrar(self, factura, estado, detalle="", valor="", evidencia=""):
        """Registra el resultado de una factura/ítem procesado."""
        self._filas.append(
            {
                "fecha_hora": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "entidad": self.entidad,
                "proceso": self.proceso,
                "factura": str(factura),
                "estado": str(estado).upper(),
                "detalle": str(detalle),
                "valor": valor,
                "evidencia": str(evidencia),
            }
        )

    def excluir(self, factura, motivo, detalle=""):
        """Registra un excluido con su motivo (va a REVISAR.csv, nunca se pierde)."""
        self._revisar.append(
            {
                "fecha_hora": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "entidad": self.entidad,
                "proceso": self.proceso,
                "factura": str(factura),
                "motivo": str(motivo),
                "detalle": str(detalle),
            }
        )

    # ------------------------------------------------------------- salidas

    def _escribir_csv(self, nombre, columnas, filas):
        ruta = os.path.join(self.carpeta, nombre)
        # utf-8-sig para que Excel en Windows abra las tildes sin problema
        with open(ruta, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=columnas, delimiter=";")
            w.writeheader()
            for fila in filas:
                w.writerow(fila)
        return ruta

    def compilar_evidencias_word(self, carpeta_imagenes, titulo=None):
        """Une los pantallazos (png/jpg) en un Word: una imagen por página.

        Requiere `python-docx`. Si no está instalado, lo informa y sigue.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            self.log("  [!] python-docx no está instalado: se omite el Word de evidencias.")
            self.log("      Instálelo con:  pip install python-docx")
            return None

        imagenes = []
        for raiz, _, archivos in os.walk(carpeta_imagenes):
            for a in sorted(archivos):
                if a.lower().endswith((".png", ".jpg", ".jpeg")):
                    imagenes.append(os.path.join(raiz, a))
        if not imagenes:
            self.log("  [!] No se hallaron imágenes en: %s" % carpeta_imagenes)
            return None

        doc = Document()
        seccion = doc.sections[0]
        seccion.page_width, seccion.page_height = Inches(8.5), Inches(11)
        ancho_util = Inches(6.5)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo or "Evidencias — %s / %s" % (self.entidad, self.proceso))
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph(
            "Generado: %s  ·  Total evidencias: %d"
            % (datetime.now().strftime("%Y-%m-%d %H:%M"), len(imagenes))
        )

        for i, img in enumerate(imagenes):
            doc.add_page_break()
            encabezado = doc.add_paragraph()
            r = encabezado.add_run(os.path.splitext(os.path.basename(img))[0])
            r.bold = True
            try:
                doc.add_picture(img, width=ancho_util)
            except Exception as e:  # imagen dañada: se anota, no se cae el lote
                doc.add_paragraph("[Imagen no legible: %s]" % e)

        ruta = os.path.join(self.carpeta, "EVIDENCIAS.docx")
        doc.save(ruta)
        self.log("  Word de evidencias: %s (%d imágenes)" % (ruta, len(imagenes)))
        return ruta

    def cerrar(self):
        """Escribe CONTROL.csv, REVISAR.csv y RESUMEN.txt. Devuelve la carpeta."""
        self._escribir_csv("CONTROL.csv", COLUMNAS_CONTROL, self._filas)
        if self._revisar:
            self._escribir_csv("REVISAR.csv", COLUMNAS_REVISAR, self._revisar)

        estados = {}
        for f in self._filas:
            estados[f["estado"]] = estados.get(f["estado"], 0) + 1
        duracion = (datetime.now() - self._inicio).total_seconds()

        lineas = [
            "RESUMEN — %s / %s" % (self.entidad, self.proceso),
            "Fecha: %s" % datetime.now().strftime("%Y-%m-%d %H:%M"),
            "Duración: %.1f s" % duracion,
            "Procesadas: %d" % len(self._filas),
        ]
        for estado, n in sorted(estados.items()):
            lineas.append("  · %s: %d" % (estado, n))
        lineas.append("Excluidas a REVISAR: %d" % len(self._revisar))

        with open(os.path.join(self.carpeta, "RESUMEN.txt"), "w", encoding="utf-8") as f:
            f.write("\n".join(lineas) + "\n")

        for l in lineas:
            self.log("  " + l)
        self.log("  Carpeta de salida: %s" % self.carpeta)
        return self.carpeta
