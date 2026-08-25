"""
nucleo.ai_tools — Funciones "inteligentes" de la caja PDF (fase 3).

Resumir, traducir, pasar a Markdown y extraer texto (OCR) de un PDF usando
Google Gemini, el MISMO proveedor que ya usa el motor de glosas
(app/services/gemini_service.py). Reutiliza la variable GEMINI_API_KEY y el
modelo GEMINI_MODEL (por defecto gemini-2.0-flash).

A diferencia del resto de la Suite, estas funciones SÍ necesitan internet y la
API key. Si no está la key, avisan con un mensaje claro y no rompen nada.

Diseño:
  · Gemini procesa PDFs de forma nativa (inline_data). El resumen manda el PDF
    entero en una sola llamada (la salida es corta).
  · Traducir / Markdown / OCR recorren PÁGINA POR PÁGINA para no truncar
    documentos largos: si la página tiene texto se envía el texto; si es un
    escaneo (sin capa de texto) se envía la imagen y Gemini la lee (OCR real).
  · La API key viaja SIEMPRE por header (x-goog-api-key), nunca en la URL
    (lección aprendida en producción del motor: la URL se loguea).

Sin dependencias nuevas: usa urllib de la librería estándar.
"""

import base64
import json
import os
import urllib.error
import urllib.request

BASE_URL = "https://generativelanguage.googleapis.com/v1beta"
DEFAULT_MODEL = os.getenv("GEMINI_MODEL") or "gemini-2.0-flash"
MAX_PAGINAS = 40  # tope de páginas por documento (coste/tiempo)
MAX_PDF_BYTES = 30 * 1024 * 1024

MENSAJE_SIN_KEY = (
    "Estas funciones usan IA (Google Gemini) y necesitan INTERNET y la "
    "variable de entorno GEMINI_API_KEY (la misma del motor de glosas). "
    "Pídala al área de sistemas y configúrela antes de usarlas. El resto de "
    "la caja PDF funciona sin esto."
)


def api_key():
    return os.getenv("GEMINI_API_KEY", "")


def disponible():
    return bool(api_key())


# ------------------------------------------------------------- transporte


def _llamar(system, parts, temperature=0.2, max_tokens=8192, model=None, timeout=120):
    """Una llamada a Gemini generateContent. Devuelve el texto de la respuesta."""
    key = api_key()
    if not key:
        raise RuntimeError(MENSAJE_SIN_KEY)
    model = model or DEFAULT_MODEL
    url = "%s/models/%s:generateContent" % (BASE_URL, model)
    body = {
        "contents": [{"role": "user", "parts": parts}],
        "systemInstruction": {"parts": [{"text": system}]},
        "generationConfig": {
            "temperature": temperature,
            "maxOutputTokens": max_tokens,
            "topP": 0.95,
            "thinkingConfig": {"thinkingBudget": 0},
        },
        "safetySettings": [
            {"category": c, "threshold": "BLOCK_NONE"}
            for c in (
                "HARM_CATEGORY_HARASSMENT",
                "HARM_CATEGORY_HATE_SPEECH",
                "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                "HARM_CATEGORY_DANGEROUS_CONTENT",
            )
        ],
    }
    datos = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=datos,
        method="POST",
        headers={"Content-Type": "application/json", "x-goog-api-key": key},
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            data = json.loads(r.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detalle = e.read()[:300].decode("utf-8", "ignore") if hasattr(e, "read") else ""
        raise RuntimeError("Gemini respondió HTTP %s: %s" % (e.code, detalle))
    except urllib.error.URLError as e:
        raise RuntimeError("No se pudo conectar con Gemini (¿hay internet?): %s" % e.reason)

    cand = (data.get("candidates") or [{}])[0]
    texto = "".join(p.get("text", "") for p in cand.get("content", {}).get("parts", [])).strip()
    if not texto:
        raise RuntimeError("Gemini no devolvió texto (motivo=%s)" % cand.get("finishReason", "?"))
    return texto


def _parte_texto(t):
    return {"text": t}


def _parte_pdf(data):
    return {
        "inline_data": {
            "mime_type": "application/pdf",
            "data": base64.standard_b64encode(data).decode("ascii"),
        }
    }


def _parte_imagen(png):
    return {
        "inline_data": {
            "mime_type": "image/png",
            "data": base64.standard_b64encode(png).decode("ascii"),
        }
    }


def _salida(ruta, sufijo):
    base = os.path.splitext(os.path.basename(ruta))[0]
    carpeta = os.path.dirname(os.path.abspath(ruta))
    return os.path.join(carpeta, base + sufijo)


def _paginas_como_partes(ruta, log):
    """Genera (n_pagina, parte, tipo) por página: texto si lo hay, si no imagen."""
    import fitz

    doc = fitz.open(ruta)
    total = doc.page_count
    n = min(total, MAX_PAGINAS)
    if total > MAX_PAGINAS:
        log("[!] El PDF tiene %d páginas; se procesan las primeras %d." % (total, MAX_PAGINAS))
    for i in range(n):
        pagina = doc[i]
        txt = pagina.get_text().strip()
        if len(txt) >= 20:
            yield i + 1, _parte_texto(txt), "texto"
        else:
            png = pagina.get_pixmap(dpi=150).tobytes("png")
            yield i + 1, _parte_imagen(png), "imagen (OCR)"
    doc.close()


# ------------------------------------------------------------- operaciones


def resumir(ruta, salida=None, log=print):
    """Resumen ejecutivo del PDF (una sola llamada; manda el PDF entero)."""
    with open(ruta, "rb") as f:
        data = f.read()
    if len(data) > MAX_PDF_BYTES:
        raise RuntimeError("El PDF pesa más de 30 MB: es demasiado para el resumen por IA.")
    system = (
        "Eres un auditor de cuentas médicas del Hospital Universitario de "
        "Santander. Resumes documentos en español claro y conciso."
    )
    user = (
        "Resume este PDF para un analista de cartera. Incluye: de qué trata, "
        "entidades/EPS, facturas o cifras clave, fechas y una conclusión. "
        "Usa viñetas y negritas en Markdown."
    )
    log("Enviando el PDF a Gemini para resumen…")
    texto = _llamar(system, [_parte_pdf(data), _parte_texto(user)], max_tokens=2048)
    salida = salida or _salida(ruta, "_resumen.md")
    with open(salida, "w", encoding="utf-8") as f:
        f.write(texto)
    log("Resumen IA → %s" % salida)
    return salida


def traducir(ruta, idioma="inglés", salida=None, log=print):
    """Traduce el PDF al idioma indicado (página por página). Devuelve un .docx."""
    from docx import Document

    system = (
        "Eres un traductor profesional. Traduces fielmente al %s conservando "
        "cifras, nombres propios, códigos y la estructura. Devuelve SOLO la "
        "traducción, sin comentarios ni notas." % idioma
    )
    doc = Document()
    doc.add_heading("Traducción al %s" % idioma, level=1)
    for npag, parte, tipo in _paginas_como_partes(ruta, log):
        log("  traduciendo página %d (%s)…" % (npag, tipo))
        pedido = _parte_texto("Traduce al %s el contenido de esta página." % idioma)
        trad = _llamar(system, [parte, pedido])
        doc.add_heading("Página %d" % npag, level=2)
        for parrafo in trad.split("\n\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())
    salida = salida or _salida(ruta, "_traducido.docx")
    doc.save(salida)
    log("Traducción IA → %s" % salida)
    return salida


def a_markdown(ruta, salida=None, log=print):
    """Convierte el PDF a Markdown limpio (página por página)."""
    system = (
        "Conviertes documentos a Markdown limpio y fiel: títulos con #, listas, "
        "y tablas en formato Markdown. Devuelve SOLO el Markdown, sin explicaciones."
    )
    trozos = []
    for npag, parte, tipo in _paginas_como_partes(ruta, log):
        log("  convirtiendo página %d (%s)…" % (npag, tipo))
        pedido = _parte_texto("Convierte a Markdown el contenido de esta página.")
        trozos.append(_llamar(system, [parte, pedido]))
    salida = salida or _salida(ruta, ".md")
    with open(salida, "w", encoding="utf-8") as f:
        f.write("\n\n".join(trozos))
    log("Markdown IA → %s" % salida)
    return salida


def ocr(ruta, salida=None, log=print):
    """Extrae TODO el texto del PDF (incluye escaneos, vía visión). Devuelve .txt."""
    system = (
        "Transcribes texto de documentos con máxima fidelidad, respetando el "
        "orden de lectura. Devuelve SOLO el texto, sin comentarios."
    )
    trozos = []
    for npag, parte, tipo in _paginas_como_partes(ruta, log):
        log("  leyendo página %d (%s)…" % (npag, tipo))
        pedido = _parte_texto("Transcribe fielmente todo el texto de esta página.")
        trozos.append(_llamar(system, [parte, pedido]))
    salida = salida or _salida(ruta, "_texto.txt")
    with open(salida, "w", encoding="utf-8") as f:
        f.write("\n\n".join(trozos))
    log("Texto extraído (OCR IA) → %s" % salida)
    return salida
