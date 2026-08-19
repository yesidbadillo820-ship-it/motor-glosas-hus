"""preauditar_glosas_adres.py — deja la macro de respuesta casi lista.

La macro `NUEVO MODELO MACRO PARA DAR RESPUESTA A GLOSA ADRES` es el reporte del
ADRES (16 columnas) **más 10 columnas que el equipo llena a mano**:

    17 CODIGO NUMERICO      21 RTA GLOSA COMPLETA      25 GESTOR
    18 CLASIFICACION        22 CANTIDAD ACEPTADA       26 MEDICO
    19 OBSERVACION (SE ACEPTA / SE OBJETA / SE SUBSANA)
    20 OBSERVACION TECNICO  23 VALOR ACEPTADO          24 CENTRO DE COSTOS

De esas diez, **siete son mecánicas** y este script las llena solo:

  - **CODIGO NUMERICO**: los primeros 4 caracteres de la causal, igual que la
    fórmula `=MID(O3,1,4)` de la macro.
  - **CLASIFICACION**: sale de la causal. La tabla (48 causales) está sacada del
    trabajo ya hecho por el equipo, y con `--macro` se vuelve a aprender del
    archivo real por si cambian criterios.
  - **CENTRO DE COSTOS**: se propone por el servicio glosado contra el catálogo
    de áreas del hospital.
  - **GESTOR / MEDICO**: del reparto por factura.
  - **RTA GLOSA COMPLETA**: réplica exacta de la fórmula de la macro.

Las tres que **no** son mecánicas —aceptar, objetar o subsanar, con su
observación y su valor aceptado— **las decide el auditor**. El bot solo
**sugiere** y explica por qué, en columnas aparte (27 en adelante) para no
mover de sitio nada de lo que usa la macro. La columna 19 solo se toca si se
pide con `--aplicar-sugerencias`.

USO (Windows, desde C:\\temp-notas):

    REM Llenar lo mecánico y sugerir el resto:
    py tools\\preauditar_glosas_adres.py ^
        --reporte-glosas "D:\\...\\ReporteGlosasReclamPAQUETE 31068.xlsx" ^
        --macro          "D:\\...\\NUEVO MODELO MACRO ... 31068.xlsm" ^
        --bitacora       "D:\\...\\BITACORA_31068.csv" ^
        --salida         "D:\\...\\MACRO_31068_PREAUDITADA.xlsx" ^
        --respuestas     "D:\\...\\RESPUESTAS"

`--macro` es opcional pero conviene: de ahí salen el catálogo de centros de
costos, el reparto gestor/médico y —a medida que el equipo vaya decidiendo— el
criterio real por causal, que el bot empieza a proponer con confianza ALTA.

Seguro por defecto: NO modifica el archivo de la macro. Escribe uno nuevo.
Requiere: py -m pip install openpyxl   (para el Word: python-docx)
"""

from __future__ import annotations

import argparse
import csv
import logging
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from ajustar_detallado_glosas import (  # noqa: E402
    _abrir_libro,
    _fila_encabezado,
    _leer_filas,
    _mapear_columnas,
    _norm,
    _norm_desc,
    normalizar_factura,
)

logger = logging.getLogger("preauditar_adres")

# ─── Estructura de la macro ──────────────────────────────────────────────────

# Las 26 columnas, en el orden exacto que espera el VBA (que trabaja por índice:
# C=3 factura, I=9 glosado, J=10 documento, N=14 elemento, U=21 respuesta,
# W=23 valor aceptado). Cambiar este orden rompe las macros.
COLUMNAS_MACRO = [
    "Código Habilitación",
    "Número Radicación",
    "Número Factura",
    "Número Paquete",
    "Cantidad Reclamado",
    "Valor Reclamado",
    "Cantidad Aprobada",
    "Valor Aprobado",
    "Valor Glosado",
    "Tip- Num Doc Victima",
    "Consecutivo Item",
    "Tipo Elemento",
    "Cod Elemento",
    "Descripción Elemento",
    "Descripción Glosa",
    "Descripción Anotación",
    "CODIGO NUMERICO",
    "CLASIFICACION DE LA GLOSA",
    "OBSERVACION (SE ACEPTA - SE OBJETA -  SE SUBSANA )",
    "OBSERVACION TECNICO / PROFESIONAL",
    "RTA GLOSA COMPLETA",
    "CANTIDAD ACEPTADA",
    "VALOR ACEPTADO",
    "CENTRO DE COSTOS",
    "GESTOR",
    "MEDICO",
]
# Columnas de apoyo del bot. Van DESPUÉS de la 26 para no correr nada.
COLUMNAS_APOYO = [
    "SUGERENCIA DEL BOT",
    "CONFIANZA",
    "POR QUÉ",
    "ESTADO EN EL DETALLADO",
    "CAUSAL (TEXTO)",
]

DECISIONES = ("SE ACEPTA", "SE OBJETA", "SE SUBSANA")

# Clasificación por causal, sacada del trabajo real del equipo sobre el paquete
# 31068 (4.619 filas). Con --macro se vuelve a aprender del archivo.
CLASIFICACION_POR_CAUSAL: dict[str, str] = {
    "1002": "GLOSADA TOTAL POR  FURIPS 1",
    "2101": "GLOSADA TOTAL POR  FURIPS 1",
    "2102": "GLOSADA TOTAL POR  FURIPS 1",
    "2103": "GLOSADA TOTAL POR  FURIPS 1",
    "2108": "GLOSADA TOTAL POR  FURIPS 2",
    "2110": "GLOSADA TOTAL POR  FURIPS 2",
    "3001": "GLOSADA TOTAL POR  FURIPS 1",
    "3102": "SOPORTES",
    "3103": "SOPORTES",
    "3104": "SOPORTES",
    "3105": "SOPORTES",
    "3106": "SOPORTES",
    "3107": "SOPORTES",
    "3108": "SOPORTES",
    "3109": "SOPORTES",
    "3110": "SOPORTES",
    "3111": "GLOSADA TOTAL POR  FURIPS 1",
    "3112": "SOPORTES",
    "3122": "GLOSADA TOTAL POR  FURIPS 1",
    "3201": "PERTINENCIA",
    "3202": "PERTINENCIA",
    "3203": "PERTINENCIA",
    "3204": "PERTINENCIA",
    "3205": "PERTINENCIA",
    "3206": "PERTINENCIA",
    "3207": "PERTINENCIA",
    "3208": "PERTINENCIA",
    "3209": "PERTINENCIA",
    "3210": "PERTINENCIA",
    "3302": "HABILITACION REPS",
    "4103": "GLOSADA TOTAL POR  FURIPS 1",
    "4104": "PERTINENCIA (NO DERIVADOS DEL SINIESTRO)",
    "4303": "TARIFAS",
    "4305": "TARIFAS",
    "4306": "TARIFAS",
    "4307": "TARIFAS",
    "4308": "TARIFAS",
    "4309": "TARIFAS",
    "4404": "CODIGOS CUPS",
    "4405": "FACOSTE",
    "4406": "TARIFAS",
    "4502": "PERTINENCIA",
    "4503": "PERTINENCIA",
    "4505": "PERTINENCIA",
    "4506": "FACTURACION",
    "4507": "PERTINENCIA",
    "4508": "PERTINENCIA",
    "4509": "FACTURACION",
}
# Sin causal, el ADRES glosó toda la reclamación por el FURIPS.
CLASIFICACION_SIN_CAUSAL = "GLOSADA TOTAL POR  FURIPS"

# Causales que trabajan DOS áreas a la vez. No es un error de criterio del
# equipo: los gestores las ven por FACTURACION y las médicas por PERTINENCIA, y
# quién la toma depende de QUÉ fue lo glosado. Por eso el sistema no las
# clasifica solo — las marca para que un SUPER ADMIN reparta.
CAUSALES_DE_DOS_AREAS: dict[str, tuple[str, str]] = {
    # 4506 — el material hace parte de otro servicio. En el paquete 31068 el
    # equipo la usó 231 veces como FACTURACION (gestores) y 24 como
    # PERTINENCIA (médicas).
    "4506": ("FACTURACION", "PERTINENCIA"),
}

# Lo que hace que una de estas glosas sea de las médicas y no de los gestores:
# material de osteosíntesis y, en general, insumos de alto costo — curación
# avanzada e instrumental de quirófano. Todo lo demás lo revisa el gestor por
# facturación.
#
# Los patrones son **específicos a propósito**: el apósito de hidrofibra con
# plata es curación avanzada, pero el apósito de gasa corriente no; la cuchilla
# para corte de hueso es de quirófano, pero la cuchilla de bisturí no. Medido
# contra las 255 filas de causal 4506 que el equipo clasificó a mano en el
# paquete 31068: coincide en 249 (97,6 %).
PISTAS_ALTO_COSTO = (
    # Implantes y osteosíntesis
    "OSTEOSINTESIS|OSTESINTESIS|PROTESIS|ENDOPROTESIS|CEMENTO OSEO|INJERTO|STENT|MARCAPASO"
    "|VALVULA CARDIACA|NEUROESTIMULADOR|DERIVACION VENTRICULO|CLIP DE ANEURISMA|ESPACIADOR"
    "|COPA ACETABULAR|SISTEMA DE FIJACION|FIJADOR EXTERNO|ALTO COSTO"
    # Curación avanzada (no la gasa ni el apósito corriente)
    "|HIDROFIBRA|HIDROCOLOIDE|AQUACEL|DUODERM|CUTICELL|GASA DE PARAFINA|CON PLATA|DE PLATA"
    "|HEMOSTATICO|CELULOSA OXIDADA|RIGENASE|FITOSTIMOLINE"
    # Instrumental de quirófano (no el bisturí corriente)
    "|FRESA|PERFORADOR CRANEAL|DERMATOMO|ARTROSCOPIA|CORTE DE HUESO|SIERRA OSCILANTE"
    "|MATERIALES DE SUTURA Y CURACION|MATERIAL DE SUTURA Y CURACION"
)

# Qué proponer según la familia de la glosa. Nada de esto se da por decidido:
# es una propuesta con su motivo, para que el auditor la confirme o la cambie.
SUGERENCIA_POR_CLASIFICACION: dict[str, tuple[str, str]] = {
    "SOPORTES": (
        "SE SUBSANA",
        "glosa por soporte ausente o incompleto: se subsana anexando el soporte",
    ),
    "TARIFAS": (
        "SE OBJETA",
        "glosa de tarifa: verificar contra el valor pactado en el contrato",
    ),
    "GLOSADA TOTAL POR  FURIPS": (
        "SE SUBSANA",
        "glosa a toda la reclamación por el FURIPS: se subsana corrigiendo o anexando el formulario",
    ),
    "GLOSADA TOTAL POR  FURIPS 1": (
        "SE SUBSANA",
        "glosa a toda la reclamación por el FURIPS: se subsana corrigiendo o anexando el formulario",
    ),
    "GLOSADA TOTAL POR  FURIPS 2": (
        "SE SUBSANA",
        "glosa a toda la reclamación por el FURIPS: se subsana corrigiendo o anexando el formulario",
    ),
}
# Familias donde el bot NO propone: las decide una persona.
SIN_SUGERENCIA: dict[str, str] = {
    "PERTINENCIA": "requiere concepto del médico auditor",
    "PERTINENCIA (NO DERIVADOS DEL SINIESTRO)": "requiere concepto del médico auditor",
    "FACTURACION": "requiere revisar la facturación del ítem",
    "CODIGOS CUPS": "requiere revisar la homologación del CUPS",
    "FACOSTE": "requiere revisar el costo del ítem",
    "HABILITACION REPS": "requiere verificar la habilitación en el REPS",
}

# Catálogo oficial de centros de costos del hospital, tal como está en la hoja
# oculta de la macro (el botón que usa el equipo). Es la lista buena: el bot no
# se inventa centros, solo escoge de acá. Si la macro que se cargue trae una
# lista distinta, esa manda (puede haber cambiado el plan de cuentas).
CATALOGO_CENTROS_COSTOS: tuple[str, ...] = (
    "510204-COSTOS",
    "510205-FACTURACION Y LIQUIDACION",
    "510401-DIREC SUBGERENCIA MUJER E INFANCIA",
    "510402-DIREC SUBGCIA SS AMBULATORIOS Y APOYO TERAPEUTICO",
    "510403-DIREC SUBGCIA MEDICAS",
    "510404-DIREC SUBGIA DE APOYO DIAGNOSTICO",
    "510405-DIREC SUBGCIA QUIRURGICA",
    "510406-DIREC SUBGCIA DE ALTO COSTO",
    "510407-DIREC SUBGCIA DE ENFERMERÍA",
    "580302-AUXILIATURA DE ENFERMERÍA",
    "580501-SERVICIO FARMACEUTICO",
    "730101-URGENCIAS PEDIATRICAS",
    "730102-URGENCIAS ADULTOS",
    "730104-URGENCIAS HOSPITALIZACION",
    "731101-CONSULTA EXTERNA ESPECIALIZADA",
    "732001-HOSPITALIZACION GINECOOBSTETRICIA",
    "732002-HOSPITALIZACION PEDIATRIA",
    "732003-HOSPITALIZACION MEDICINA INTERNA",
    "732005-HOSPITALIZACION ESPECIALIDADES QUIRURGICAS",
    "732006-HOSPITALIZACION HEMATOLOGIA",
    "732007-HOSPITALIZACION CIRUGIA GENERAL",
    "732101-UCI ADULTOS - MEDICAS P-1",
    "732102-UCI PEDIATRICA",
    "732104-UCI NEONATAL P4",
    "732107-UCI ADULTOS ALTO COSTO",
    "732108-UCI ADULTOS ALTO COSTO P-4",
    "732109-UCI ADULTOS - MEDICAS",
    "732301-RECIEN NACIDOS",
    "732501-QUEMADOS",
    "733001-QUIROFANOS",
    "733101-URGENCIAS GINECOBSTETRICAS",
    "734001-LABORATORIO CLINICO",
    "734101-RADIOLOGIA",
    "734102-ECOGRAFIA APOYO DIAGNOSTICO",
    "734103-ESCANOGRAFIA",
    "734106-ANGIOGRAFIA",
    "734107-EQUIPO PORTATIL IMAGENES",
    "734201-PATOLOGIA",
    "734301-ECOGRAFIA GINECOBSTETRICA",
    "734302-GASTROENTEROLOGIA",
    "734303-CARDIOLOGIA Y HERMODINAMIA",
    "734304-ACCESOS VASCULARES",
    "734901-FISIOTERAPIA Y REHABILITACION",
    "735101-BANCO DE SANGRE",
    "735201-UNIDAD RENAL",
)

# Índice por nombre (sin el código) para poder devolver siempre la forma
# oficial "código-NOMBRE" aunque internamente se razone con el nombre.
_CENTRO_POR_NOMBRE: dict[str, str] = {
    entrada.split("-", 1)[1].strip().upper(): entrada for entrada in CATALOGO_CENTROS_COSTOS
}


def centro_oficial(nombre: str, catalogo: tuple[str, ...] | list[str] | None = None) -> str:
    """Devuelve el centro en su forma oficial `código-NOMBRE`.

    Acepta que le pasen el nombre solo ("QUIROFANOS"), la forma completa o algo
    que no está en el catálogo — en ese último caso devuelve lo que le dieron,
    para no perder lo que el equipo escribió a mano.
    """
    texto = (nombre or "").strip()
    if not texto:
        return ""
    indice = (
        {e.split("-", 1)[1].strip().upper(): e for e in catalogo if "-" in e}
        if catalogo
        else _CENTRO_POR_NOMBRE
    )
    if texto.upper() in {e.upper() for e in (catalogo or CATALOGO_CENTROS_COSTOS)}:
        return texto
    sin_codigo = texto.split("-", 1)[1].strip() if re.match(r"^\d{6}-", texto) else texto
    return indice.get(sin_codigo.upper(), texto)


# Palabras del servicio glosado → área del hospital. Se recorre en orden.
PISTAS_CENTRO_COSTOS: tuple[tuple[str, str], ...] = (
    (
        "BANCO DE SANGRE",
        "UNIDAD DE SANGRE|GLOBULOS ROJOS|HEMOCLASIFICACION|COMPATIBILIDAD"
        "|ANTICUERPOS IRREGULARES|PLASMA|PLAQUETAS|CRIOPRECIPITADO",
    ),
    (
        "LABORATORIO CLINICO",
        "HEMOGRAMA|CUADRO HEMATICO|GLUCOSA|GLUCOMETRIA|CREATININA|NITROGENO UREICO"
        "|TRANSAMINASA|PROTROMBINA|TROMBOPLASTINA|CULTIVO|ANTIBIOGRAMA|PROTEINA C REACTIVA"
        "|SEDIMENTACION|UROANALISIS|SODIO|POTASIO|CLORURO|CLORO|BILIRRUBINA|AMILASA"
        "|PROCALCITONINA|GASES ARTERIALES|IONOGRAMA|PARCIAL DE ORINA|CALCITONINA"
        "|COOMBS|LACTATO|CALCIO COLORIMETRICO|MAGNESIO|CREATINCINASA|FOSFATASA|ALBUMINA"
        "|SUERO ORINA Y OTROS|FACTOR RH|PRUEBA CUALITATIVA|PRUEBA CUANTITATIVA|TROPONINA"
        "|DIMERO|FERRITINA|TSH|HORMONA|ANTIGENO|SEROLOGIA|PARCIAL DE SANGRE",
    ),
    ("ESCANOGRAFIA", "TOMOGRAFIA|ESCANOGRAFIA|TAC |RECONSTRUCCION TRIDIMENSIONAL|RESONANCIA"),
    ("ECOGRAFIA APOYO DIAGNOSTICO", "ECOGRAFIA|DOPPLER|DUPLEX"),
    (
        "RADIOLOGIA",
        "RADIOGRAFIA|RAYOS X|TORAX|COLUMNA|CRANEO|ABDOMEN TOTAL|CADERA"
        "|EXTREMIDADES Y ARTICULACIONES|BRAZO PIERNA|MANO DEDOS|SENOS PARANASALES"
        "|PROYECCION ADICIONAL|COMPARATIVAS DE LAS REGIONES|RIÑONES BAZO AORTA"
        "|PELVIS|CLAVICULA|ANTEBRAZO|OMOPLATO",
    ),
    ("EQUIPO PORTATIL IMAGENES", "PORTATIL|FLUOROSCOPIA|INTENSIFICADOR"),
    ("PATOLOGIA", "PATOLOGIA|BIOPSIA|CITOLOGIA"),
    (
        "FISIOTERAPIA Y REHABILITACION",
        "TERAPIA FISICA|TERAPIA OCUPACIONAL|TERAPIA DEL LENGUAJE"
        "|TERAPIA RESPIRATORIA|REHABILITACION|FONOAUDIOLOGIA",
    ),
    (
        "QUIROFANOS",
        "DERECHOS DE SALA DE CIRUGIA|HONORARIOS CIRUJANO|SERVICIOS PROFESIONALES DEL CIRUJANO"
        "|ANESTESIOLOGO|AYUDANTIA QUIRURGICA|MATERIALES DE SUTURA|MATERIALES Y SUTURAS|OSTEOSINTESIS",
    ),
    ("URGENCIAS ADULTOS", "CONSULTA DE URGENCIAS|DERECHOS DE SALA DE YESOS|INMOVILIZACION"),
    ("CONSULTA EXTERNA ESPECIALIZADA", "CONSULTA AMBULATORIA|INTERCONSULTA|VALORACION INICIAL"),
    ("UNIDAD RENAL", "DIALISIS|HEMODIALISIS|UNIDAD RENAL"),
    ("SERVICIO FARMACEUTICO", "MEDICAMENTO"),  # respaldo por Tipo Elemento
)
CENTRO_POR_TIPO_ELEMENTO: dict[str, str] = {
    "MEDICAMENTOS": "SERVICIO FARMACEUTICO",
    "DISPOSITIVOS MEDICOS": "SERVICIO FARMACEUTICO",
    "INSUMO": "SERVICIO FARMACEUTICO",
    "MATERIAL DE OSTEOSINTESIS": "QUIROFANOS",
}
# Servicios que NO se pueden asignar a un área sin más contexto: la habitación
# no dice de qué especialidad es, y las glosas a la reclamación no son de un
# servicio. Se dejan en blanco a propósito, no por olvido.
SIN_CENTRO_CLARO = (
    "HABITACION|SALA ESPECIAL|ATENCION DIARIA INTRAHOSPITALARIA|MATERIALES DE CURACION"
)

TEXTO_EXTEMPORANEA = (
    "ADICIONALMENTE SE INFORMA QUE LA INCONFORMIDAD NO SE RADICO EN LOS TIEMPOS "
    "ESTABLECIDOS SEGÚN RESOLUCION 1236 DE 2023 EN SU ARTICULO 8 NUMERAL 8.5, POR LO "
    "TANTO, CORRESPONDE A UNA GLOSA EXTEMPORANEA."
)
ENCABEZADO_RESPUESTA = (
    "SE IDENTIFICAN HALLAZGOS Y SE REALIZAN LOS AJUSTES CORRESPONDIENTES POR GLOSA "
    "ACEPTADA PARCIAL POR VALOR DE ${valor}"
)

_ALIAS_EXTRA = {
    "cod_habilitacion": ("CODIGO HABILITACION", "COD HABILITACION"),
    "doc_victima": ("TIP- NUM DOC VICTIMA", "TIP NUM DOC VICTIMA", "DOC VICTIMA"),
}


# ─── Modelo ──────────────────────────────────────────────────────────────────


@dataclass
class FilaMacro:
    """Una fila del reporte con lo que el bot le agrega."""

    # Del ADRES
    cod_habilitacion: str = ""
    radicacion: str = ""
    factura: str = ""
    paquete: str = ""
    cant_reclamada: float = 0.0
    valor_reclamado: float = 0.0
    cant_aprobada: float = 0.0
    valor_aprobado: float = 0.0
    valor_glosado: float = 0.0
    doc_victima: str = ""
    consecutivo: str = ""
    tipo_elemento: str = ""
    codigo: str = ""
    descripcion: str = ""
    descripcion_glosa: str = ""
    anotacion: str = ""
    # Del bot
    codigo_numerico: str = ""
    clasificacion: str = ""
    observacion: str = ""
    observacion_tecnico: str = ""
    cantidad_aceptada: str = ""
    valor_aceptado: float = 0.0
    centro_costos: str = ""
    gestor: str = ""
    medico: str = ""
    sugerencia: str = ""
    confianza: str = ""
    motivo: str = ""
    estado_detallado: str = ""

    @property
    def rta_glosa_completa(self) -> str:
        return rta_glosa_completa(self)


@dataclass
class Resumen:
    filas: list[FilaMacro] = field(default_factory=list)
    aprendidas: dict[str, tuple[str, int, int]] = field(default_factory=dict)

    @property
    def facturas(self) -> int:
        return len({normalizar_factura(f.factura) for f in self.filas})


# ─── Reglas ──────────────────────────────────────────────────────────────────


def codigo_numerico(descripcion_glosa: str) -> str:
    """Los primeros 4 caracteres de la causal, igual que `=MID(O3,1,4)`."""
    return (descripcion_glosa or "")[:4].strip()


def clasificar(causal: str, tabla: dict[str, str] | None = None) -> str:
    tabla = tabla if tabla is not None else CLASIFICACION_POR_CAUSAL
    if not causal:
        return CLASIFICACION_SIN_CAUSAL
    return tabla.get(causal, CLASIFICACION_POR_CAUSAL.get(causal, ""))


def necesita_asignacion(causal: str) -> bool:
    """¿Esta causal la trabajan dos áreas y hay que repartirla a mano?"""
    return causal in CAUSALES_DE_DOS_AREAS


def reparto_de_area(causal: str, tipo_elemento: str, descripcion: str) -> tuple[str, str]:
    """Para las causales de dos áreas, a quién le debería tocar.

    Devuelve (área sugerida, motivo). **No decide**: el SUPER ADMIN confirma,
    porque depende del procedimiento y de lo que se glosó.
    """
    opciones = CAUSALES_DE_DOS_AREAS.get(causal)
    if not opciones:
        return "", ""
    facturacion, pertinencia = opciones
    texto = _norm_desc(descripcion)
    if _norm(tipo_elemento) == "MATERIAL DE OSTEOSINTESIS" or re.search(PISTAS_ALTO_COSTO, texto):
        return pertinencia, (
            "material de osteosíntesis o de alto costo: la revisa el médico auditor"
        )
    return facturacion, (
        "no parece material de osteosíntesis ni de alto costo: la revisa el gestor por facturación"
    )


def centro_de_costos(
    tipo_elemento: str, descripcion: str, catalogo: tuple[str, ...] | list[str] | None = None
) -> str:
    """Área del hospital que corresponde al servicio glosado (propuesta).

    Sale siempre en la forma oficial `código-NOMBRE` del catálogo del hospital.
    """
    texto = _norm_desc(descripcion)
    if not texto or re.search(SIN_CENTRO_CLARO, texto):
        return ""
    for area, patron in PISTAS_CENTRO_COSTOS:
        if re.search(patron, texto):
            return centro_oficial(area, catalogo)
    respaldo = CENTRO_POR_TIPO_ELEMENTO.get(_norm(tipo_elemento), "")
    return centro_oficial(respaldo, catalogo) if respaldo else ""


def sugerir(
    clasificacion: str, causal: str, aprendidas: dict[str, tuple[str, int, int]] | None = None
) -> tuple[str, str, str]:
    """(sugerencia, confianza, motivo).

    Si el equipo ya decidió lo suficiente para esa causal, se propone **su**
    criterio con confianza ALTA. Si no, se cae en la regla por familia, con
    confianza MEDIA. Y donde hace falta criterio humano, no se propone nada.
    """
    aprendidas = aprendidas or {}
    if causal in aprendidas:
        decision, n, total = aprendidas[causal]
        return (
            decision,
            "APRENDIDA",
            f"el equipo decidió así en {n} de {total} glosas ya resueltas con causal {causal}",
        )

    if clasificacion in SIN_SUGERENCIA:
        return "", "", SIN_SUGERENCIA[clasificacion]

    propuesta = SUGERENCIA_POR_CLASIFICACION.get(clasificacion)
    if propuesta:
        return propuesta[0], "REGLA", propuesta[1]
    return "", "", "sin regla para esta clasificación: lo define el auditor"


def _moneda(valor: float) -> str:
    """`TEXT(valor,"$#.##0")` de la macro: $31.800 (punto de miles)."""
    return "$" + f"{int(round(valor or 0)):,}".replace(",", ".")


def rta_glosa_completa(f: FilaMacro) -> str:
    """Réplica exacta de la fórmula de la columna U de la macro.

    =IF(W>0, CONCAT(Q,"-",S," ",N,"   CANTIDAD ACEPTADA ",V," . POR VALOR ",
                    TEXT(W,"$#.##0")," ",T),
             CONCAT(Q,"-",S,"-",N,"POR VALOR DE  ",TEXT(I,"$#.##0")," ",T))
    """
    if (f.valor_aceptado or 0) > 0:
        return (
            f"{f.codigo_numerico}-{f.observacion} {f.descripcion}  "
            f"CANTIDAD ACEPTADA {f.cantidad_aceptada} . POR VALOR "
            f"{_moneda(f.valor_aceptado)} {f.observacion_tecnico}"
        )
    return (
        f"{f.codigo_numerico}-{f.observacion}-{f.descripcion}POR VALOR DE  "
        f"{_moneda(f.valor_glosado)} {f.observacion_tecnico}"
    )


# ─── Lectura ─────────────────────────────────────────────────────────────────


def leer_reporte(
    ruta: Path, paquete: str | None = None, minimo_glosado: float = 1.0
) -> list[FilaMacro]:
    """Las filas GLOSADAS del reporte del ADRES, con sus 16 columnas.

    Se dejan afuera las glosadas por menos de un peso: son redondeos del ADRES
    (en el paquete 31068 son 19 filas de un mismo insumo) y la macro del equipo
    tampoco las trae.
    """
    from ajustar_detallado_glosas import _ALIAS_COL_GLOSA

    alias = {**_ALIAS_COL_GLOSA, **_ALIAS_EXTRA}
    filas = _leer_filas(ruta)
    idx = _fila_encabezado(filas, alias, minimo=4)
    if idx < 0:
        raise SystemExit(f"No reconocí el encabezado del reporte en {ruta.name}.")
    cols = _mapear_columnas(filas[idx], alias)

    def val(fila, campo, defecto=""):
        j = cols.get(campo, -1)
        v = fila[j] if 0 <= j < len(fila) else None
        return defecto if v is None else v

    from ajustar_detallado_glosas import _parse_valor

    salida: list[FilaMacro] = []
    for fila in filas[idx + 1 :]:
        factura = str(val(fila, "factura")).strip()
        if not factura:
            continue
        if paquete and _norm(val(fila, "paquete")) != _norm(paquete):
            continue
        glosado = _parse_valor(val(fila, "valor_glosado"))
        if glosado < minimo_glosado:
            continue  # la macro solo trabaja lo que sigue glosado
        salida.append(
            FilaMacro(
                cod_habilitacion=str(val(fila, "cod_habilitacion")).strip(),
                radicacion=str(val(fila, "radicacion")).strip(),
                factura=factura,
                paquete=str(val(fila, "paquete")).strip(),
                cant_reclamada=_parse_valor(val(fila, "cant_reclamada")),
                valor_reclamado=_parse_valor(val(fila, "valor_reclamado")),
                cant_aprobada=_parse_valor(val(fila, "cant_aprobada")),
                valor_aprobado=_parse_valor(val(fila, "valor_aprobado")),
                valor_glosado=glosado,
                doc_victima=str(val(fila, "doc_victima")).strip(),
                consecutivo=str(val(fila, "consecutivo")).strip(),
                tipo_elemento=str(val(fila, "tipo_elemento")).strip(),
                codigo=str(val(fila, "codigo")).strip(),
                # Sin .strip(): el ADRES cierra algunas descripciones con un
                # espacio de no separación y el texto de respuesta tiene que
                # salir carácter por carácter igual al de la macro.
                descripcion=str(val(fila, "descripcion")),
                descripcion_glosa=str(val(fila, "descripcion_glosa")),
                anotacion=str(val(fila, "anotacion")),
            )
        )
    return salida


def leer_macro(
    ruta: Path,
) -> tuple[dict[str, str], dict[str, tuple[str, str]], list[str], list[dict]]:
    """Del libro de la macro saca lo que el equipo ya construyó.

    Devuelve (clasificación por causal, reparto por factura, catálogo de centros
    de costos, filas ya trabajadas).
    """
    wb = _abrir_libro(ruta, solo_lectura=True)
    try:
        hoja = wb.worksheets[0]
        filas = []
        for fila in hoja.iter_rows(min_row=3, values_only=True):
            if not fila or fila[2] is None:
                break
            filas.append(dict(zip(COLUMNAS_MACRO, fila, strict=False)))

        clasif: dict[str, Counter] = defaultdict(Counter)
        reparto: dict[str, tuple[str, str]] = {}
        for f in filas:
            causal = codigo_numerico(str(f.get("Descripción Glosa") or ""))
            etiqueta = str(f.get("CLASIFICACION DE LA GLOSA") or "").strip()
            if causal and etiqueta:
                clasif[causal][etiqueta] += 1
            clave = normalizar_factura(str(f.get("Número Factura") or ""))
            if clave not in reparto:
                reparto[clave] = (
                    str(f.get("GESTOR") or "").strip(),
                    str(f.get("MEDICO") or "").strip(),
                )

        centros: list[str] = []
        for ws in wb.worksheets[1:]:
            valores = [str(r[0]).strip() for r in ws.iter_rows(values_only=True) if r and r[0]]
            if valores and re.match(r"^\d{6}-", valores[0]):
                centros = valores
                break
        return (
            {c: v.most_common(1)[0][0] for c, v in clasif.items()},
            reparto,
            centros,
            filas,
        )
    finally:
        wb.close()


def _clave_fila(factura: str, codigo: str, glosado) -> tuple:
    return (normalizar_factura(str(factura or "")), _norm(codigo), round(float(glosado or 0), 2))


def trabajo_ya_hecho(filas_macro: list[dict]) -> dict[tuple, list[dict]]:
    """Indexa las filas de la macro para poder devolverle al equipo lo suyo.

    La clave es factura + código del ítem + valor glosado, y se guarda la lista
    de repeticiones en orden: el reporte trae el mismo ítem varias veces y a
    cada repetición le puede corresponder un médico distinto.
    """
    idx: dict[tuple, list[dict]] = defaultdict(list)
    for f in filas_macro:
        clave = _clave_fila(f.get("Número Factura"), f.get("Cod Elemento"), f.get("Valor Glosado"))
        idx[clave].append(f)
    return idx


def aprender_decisiones(
    filas_macro: list[dict], minimo: int = 5, acuerdo: float = 0.8
) -> dict[str, tuple[str, int, int]]:
    """Criterio del equipo por causal, a partir de lo que ya decidieron.

    Solo se toma cuando hay al menos `minimo` decisiones para esa causal y al
    menos el `acuerdo` de ellas coincide. Con eso el bot deja de adivinar y
    empieza a proponer **el criterio de la casa**.
    """
    por_causal: dict[str, Counter] = defaultdict(Counter)
    for f in filas_macro:
        decision = _norm(f.get("OBSERVACION (SE ACEPTA - SE OBJETA -  SE SUBSANA )"))
        if decision not in DECISIONES:
            continue
        causal = codigo_numerico(str(f.get("Descripción Glosa") or ""))
        if causal:
            por_causal[causal][decision] += 1

    aprendidas: dict[str, tuple[str, int, int]] = {}
    for causal, cuenta in por_causal.items():
        decision, n = cuenta.most_common(1)[0]
        total = sum(cuenta.values())
        if total >= minimo and n / total >= acuerdo:
            aprendidas[causal] = (decision, n, total)
    return aprendidas


def leer_bitacora(ruta: Path) -> dict[tuple[str, str], str]:
    """Estado de cada ítem según el ajustador: (factura, código) → acción."""
    estados: dict[tuple[str, str], str] = {}
    with ruta.open(encoding="utf-8-sig", newline="") as fh:
        for fila in csv.DictReader(fh, delimiter=";"):
            accion = (fila.get("ACCION") or "").strip()
            if accion in ("ELIMINADA", ""):
                continue
            clave = (normalizar_factura(fila.get("FACTURA", "")), _norm(fila.get("CODIGO", "")))
            estados.setdefault(clave, accion)
    return estados


def leer_reparto(ruta: Path) -> dict[str, tuple[str, str]]:
    """Excel/CSV con FACTURA, GESTOR y (opcional) MEDICO."""
    filas = _leer_filas(ruta)
    alias = {
        "factura": ("FACTURA", "NUMERO FACTURA", "NRO FACTURA"),
        "gestor": ("GESTOR", "AUDITOR", "RESPONSABLE"),
        "medico": ("MEDICO", "MÉDICO", "PROFESIONAL"),
    }
    idx = _fila_encabezado(filas, alias, minimo=2)
    if idx < 0:
        raise SystemExit(f"No reconocí el encabezado del reparto en {ruta.name}.")
    cols = _mapear_columnas(filas[idx], alias)

    def val(fila, campo):
        j = cols.get(campo, -1)
        return str(fila[j]).strip() if 0 <= j < len(fila) and fila[j] is not None else ""

    salida: dict[str, tuple[str, str]] = {}
    for fila in filas[idx + 1 :]:
        clave = normalizar_factura(val(fila, "factura"))
        if clave and clave != "0":
            salida[clave] = (val(fila, "gestor"), val(fila, "medico"))
    return salida


# ─── Preauditoría ────────────────────────────────────────────────────────────


def preauditar(
    filas: list[FilaMacro],
    *,
    tabla_clasificacion: dict[str, str] | None = None,
    reparto: dict[str, tuple[str, str]] | None = None,
    aprendidas: dict[str, tuple[str, int, int]] | None = None,
    estados: dict[tuple[str, str], str] | None = None,
    hecho: dict[tuple, list[dict]] | None = None,
    aplicar: str = "no",
) -> Resumen:
    """Llena lo mecánico, respeta lo ya trabajado y anota la sugerencia."""
    reparto = reparto or {}
    estados = estados or {}
    aprendidas = aprendidas or {}
    hecho = hecho or {}
    consumidas: Counter = Counter()
    rescatadas = 0

    for f in filas:
        # Lo que una persona ya escribió manda: se copia tal cual y no se pisa.
        clave = _clave_fila(f.factura, f.codigo, f.valor_glosado)
        repeticiones = hecho.get(clave)
        if repeticiones:
            n = consumidas[clave]
            consumidas[clave] += 1
            previa = repeticiones[n] if n < len(repeticiones) else repeticiones[-1]
            # Sin .strip(): el equipo escribe "SE ACEPTA " con espacio al final,
            # y ese espacio forma parte del texto que se le manda al ADRES.
            f.observacion = str(
                previa.get("OBSERVACION (SE ACEPTA - SE OBJETA -  SE SUBSANA )") or ""
            )
            f.observacion_tecnico = str(previa.get("OBSERVACION TECNICO / PROFESIONAL") or "")
            f.cantidad_aceptada = str(previa.get("CANTIDAD ACEPTADA") or "").strip()
            f.valor_aceptado = float(previa.get("VALOR ACEPTADO") or 0)
            f.gestor = str(previa.get("GESTOR") or "").strip()
            f.medico = str(previa.get("MEDICO") or "").strip()
            centro = str(previa.get("CENTRO DE COSTOS") or "").strip()
            if centro:
                f.centro_costos = centro
            if f.observacion:
                rescatadas += 1

        f.codigo_numerico = codigo_numerico(f.descripcion_glosa)
        f.clasificacion = clasificar(f.codigo_numerico, tabla_clasificacion)
        if not f.centro_costos:
            f.centro_costos = centro_de_costos(f.tipo_elemento, f.descripcion)
        clave_factura = normalizar_factura(f.factura)
        if not f.gestor:
            f.gestor, f.medico = reparto.get(clave_factura, (f.gestor, f.medico))
        f.sugerencia, f.confianza, f.motivo = sugerir(
            f.clasificacion, f.codigo_numerico, aprendidas
        )
        f.estado_detallado = estados.get((clave_factura, _norm(f.codigo)), "")
        if (
            not f.observacion
            and f.sugerencia
            and (aplicar == "todas" or (aplicar == "aprendidas" and f.confianza == "APRENDIDA"))
        ):
            f.observacion = f.sugerencia
    if rescatadas:
        logger.info("Se conservaron %d decisión(es) que el equipo ya había escrito", rescatadas)
    return Resumen(filas=filas, aprendidas=aprendidas)


# ─── Escritura del libro ─────────────────────────────────────────────────────


def escribir_libro(resumen: Resumen, salida: Path) -> None:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Hoja1"

    # Fila 1: los totales que la macro deja arriba (I1 y W1).
    ws.cell(row=1, column=9, value=sum(f.valor_glosado for f in resumen.filas))
    ws.cell(row=1, column=23, value=sum(f.valor_aceptado for f in resumen.filas))

    encabezados = COLUMNAS_MACRO + COLUMNAS_APOYO
    gris = PatternFill("solid", fgColor="D9D9D9")
    azul = PatternFill("solid", fgColor="DDEBF7")
    for j, titulo in enumerate(encabezados, start=1):
        celda = ws.cell(row=2, column=j, value=titulo)
        celda.font = Font(bold=True)
        celda.fill = azul if j > len(COLUMNAS_MACRO) else gris
        celda.alignment = Alignment(wrap_text=True, vertical="center")

    for i, f in enumerate(resumen.filas, start=3):
        valores = [
            f.cod_habilitacion,
            f.radicacion,
            f.factura,
            f.paquete,
            f.cant_reclamada,
            f.valor_reclamado,
            f.cant_aprobada,
            f.valor_aprobado,
            f.valor_glosado,
            f.doc_victima,
            f.consecutivo,
            f.tipo_elemento,
            f.codigo,
            f.descripcion,
            f.descripcion_glosa,
            f.anotacion,
            f.codigo_numerico,
            f.clasificacion,
            f.observacion,
            f.observacion_tecnico,
            f.rta_glosa_completa,
            f.cantidad_aceptada,
            f.valor_aceptado or 0,
            f.centro_costos,
            f.gestor,
            f.medico,
            f.sugerencia,
            f.confianza,
            f.motivo,
            f.estado_detallado,
            f.descripcion_glosa,
        ]
        for j, v in enumerate(valores, start=1):
            ws.cell(row=i, column=j, value=v)

    anchos = {
        3: 14,
        12: 20,
        13: 20,
        14: 46,
        15: 40,
        16: 40,
        18: 26,
        19: 18,
        20: 26,
        21: 60,
        24: 34,
        25: 12,
        26: 12,
        27: 14,
        28: 10,
        29: 46,
        30: 18,
    }
    for col, ancho in anchos.items():
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = ancho
    ws.freeze_panes = "A3"
    ws.auto_filter.ref = (
        f"A2:{openpyxl.utils.get_column_letter(len(encabezados))}{len(resumen.filas) + 2}"
    )

    salida.parent.mkdir(parents=True, exist_ok=True)
    wb.save(salida)


# ─── Word de respuesta por factura (réplica del VBA) ─────────────────────────


def armar_texto_respuesta(filas: list[FilaMacro], prefijo: str = "") -> str:
    """Réplica de `GenerarWord_Factura_Ordenada`.

    Junta las respuestas repetidas, pone primero las **aceptadas** ordenadas por
    valor de mayor a menor, después las objetadas o subsanadas, y cierra con el
    párrafo de glosa extemporánea.
    """
    # Import acá adentro: `glosas_adres_por_factura` importa `FilaMacro` de
    # este módulo, así que arriba haría un círculo.
    from glosas_adres_por_factura import aceptado_sin_duplicar, marcar_conteo

    aceptadas: dict[str, float] = {}
    otras: list[str] = []
    # El total que se le declara al ADRES no puede contar dos veces el mismo
    # ítem. El reporte abre un renglón por causal y el gestor decide renglón
    # por renglón: aceptando en las dos causales de un TAC glosado $700.000,
    # el encabezado decía $1.400.000. 19-08-2026.
    marcas = marcar_conteo(filas)
    total_aceptado = aceptado_sin_duplicar(
        (
            (
                (f.factura or "").strip().upper(),
                (f.codigo or "").strip().upper(),
                round(f.valor_glosado or 0, 2),
            ),
            f.valor_glosado or 0,
            f.valor_aceptado or 0,
            marca,
        )
        for f, marca in zip(filas, marcas, strict=True)
    )

    for f in filas:
        texto = f.rta_glosa_completa
        arriba = texto.upper()
        if any(m in arriba for m in ("NO SE ACEPTA GLOSA", "SE SUBSANA", "SE OBJETA")):
            if texto not in otras:
                otras.append(texto)
        elif "SE ACEPTA" in arriba:
            aceptadas[texto] = aceptadas.get(texto, 0.0) + (f.valor_aceptado or 0)

    partes = [
        prefijo + ENCABEZADO_RESPUESTA.format(valor=_moneda(total_aceptado).lstrip("$")),
        "",
    ]
    partes += [t for t, _ in sorted(aceptadas.items(), key=lambda kv: -kv[1])]
    partes += otras
    partes += ["", TEXTO_EXTEMPORANEA]
    return "\n".join(partes)


def escribir_word(texto: str, ruta: Path, titulo: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return False
    doc = Document()
    encabezado = doc.add_paragraph()
    corrida = encabezado.add_run(titulo)
    corrida.bold = True
    corrida.font.size = Pt(14)
    for linea in texto.split("\n"):
        doc.add_paragraph(linea)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    doc.save(ruta)
    return True


def generar_respuestas(
    resumen: Resumen, carpeta: Path, *, prefijo: str = "", carpeta_por_factura: bool = False
) -> tuple[int, int]:
    """Un documento de respuesta por factura. (generados, en_texto_plano)."""
    por_factura: dict[str, list[FilaMacro]] = defaultdict(list)
    for f in resumen.filas:
        por_factura[f.factura].append(f)

    hechos = planos = 0
    for factura, filas in sorted(por_factura.items()):
        texto = armar_texto_respuesta(filas, prefijo=prefijo)
        base = carpeta / factura if carpeta_por_factura else carpeta
        titulo = f"RESPUESTA A GLOSA — {factura}"
        if escribir_word(texto, base / f"RTA_GLOSA_{factura}.docx", titulo):
            hechos += 1
        else:
            base.mkdir(parents=True, exist_ok=True)
            (base / f"RTA_GLOSA_{factura}.txt").write_text(
                f"{titulo}\n\n{texto}\n", encoding="utf-8"
            )
            hechos += 1
            planos += 1
    return hechos, planos


# ─── CLI ─────────────────────────────────────────────────────────────────────


def construir_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="Llena lo mecánico de la macro de respuesta y sugiere el resto.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    p.add_argument("--reporte-glosas", type=Path, required=True)
    p.add_argument("--macro", type=Path, help="Libro de la macro (catálogos, reparto y criterio).")
    p.add_argument("--bitacora", type=Path, help="Bitácora del ajustador de detallados.")
    p.add_argument("--reparto", type=Path, help="Excel/CSV con FACTURA, GESTOR y MEDICO.")
    p.add_argument("--salida", type=Path, required=True, help="Libro resultante (.xlsx).")
    p.add_argument("--respuestas", type=Path, help="Carpeta para el Word de respuesta por factura.")
    p.add_argument("--carpeta-por-factura", action="store_true")
    p.add_argument(
        "--prefijo-respuesta", default="", help="Texto antes del encabezado (ej. '123 ')."
    )
    p.add_argument("--paquete", help="Filtrar el reporte por número de paquete.")
    p.add_argument(
        "--minimo-glosado",
        type=float,
        default=1.0,
        help="Ignorar las filas glosadas por menos de este valor (redondeos del ADRES).",
    )
    p.add_argument(
        "--aplicar-sugerencias",
        choices=("no", "aprendidas", "todas"),
        default="no",
        help="Escribir la sugerencia en la columna OBSERVACION. Por defecto NO: "
        "queda al lado para que el auditor decida. 'aprendidas' aplica solo las que "
        "salen del criterio que el propio equipo ya usó.",
    )
    p.add_argument("--reporte-csv", type=Path, help="Resumen por causal de lo que hizo.")
    p.add_argument("--verbose", action="store_true")
    return p


def main(argv: list[str] | None = None) -> int:
    args = construir_parser().parse_args(argv)
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)-8s %(message)s",
    )

    if not args.reporte_glosas.exists():
        logger.error("No existe el reporte de glosas: %s", args.reporte_glosas)
        return 2
    filas = leer_reporte(
        args.reporte_glosas, paquete=args.paquete, minimo_glosado=args.minimo_glosado
    )
    logger.info("Reporte: %d fila(s) glosada(s)", len(filas))
    if not filas:
        logger.error("El reporte no trajo ninguna fila glosada.")
        return 2

    tabla = dict(CLASIFICACION_POR_CAUSAL)
    reparto: dict[str, tuple[str, str]] = {}
    aprendidas: dict[str, tuple[str, int, int]] = {}
    hecho: dict[tuple, list[dict]] = {}
    if args.macro:
        if not args.macro.exists():
            logger.error("No existe la macro: %s", args.macro)
            return 2
        aprendida_clasif, reparto, centros, filas_macro = leer_macro(args.macro)
        tabla.update(aprendida_clasif)
        aprendidas = aprender_decisiones(filas_macro)
        hecho = trabajo_ya_hecho(filas_macro)
        logger.info(
            "Macro: %d causal(es) de clasificación, %d factura(s) con reparto, "
            "%d centro(s) de costos, %d fila(s) ya trabajadas",
            len(aprendida_clasif),
            len(reparto),
            len(centros),
            len(filas_macro),
        )
        if aprendidas:
            logger.info("Criterio aprendido del equipo en %d causal(es):", len(aprendidas))
            for causal, (decision, n, total) in sorted(aprendidas.items()):
                logger.info("   %s → %s  (%d de %d)", causal, decision, n, total)
        else:
            logger.info(
                "Todavía no hay decisiones suficientes para aprender el criterio: "
                "las sugerencias van con confianza MEDIA."
            )

    if args.reparto:
        reparto.update(leer_reparto(args.reparto))
        logger.info("Reparto: %d factura(s)", len(reparto))

    estados = {}
    if args.bitacora:
        if not args.bitacora.exists():
            logger.error("No existe la bitácora: %s", args.bitacora)
            return 2
        estados = leer_bitacora(args.bitacora)
        logger.info("Bitácora: %d ítem(s) con estado en el detallado", len(estados))

    resumen = preauditar(
        filas,
        tabla_clasificacion=tabla,
        reparto=reparto,
        aprendidas=aprendidas,
        estados=estados,
        hecho=hecho,
        aplicar=args.aplicar_sugerencias,
    )

    escribir_libro(resumen, args.salida)
    logger.info("Escrito: %s", args.salida)

    if args.respuestas:
        hechos, planos = generar_respuestas(
            resumen,
            args.respuestas,
            prefijo=args.prefijo_respuesta,
            carpeta_por_factura=args.carpeta_por_factura,
        )
        logger.info("Respuestas: %d documento(s) en %s", hechos, args.respuestas)
        if planos:
            logger.warning(
                "%d salieron como .txt porque falta python-docx (py -m pip install python-docx)",
                planos,
            )

    _resumir(resumen)
    if args.reporte_csv:
        _escribir_reporte(resumen, args.reporte_csv)
        logger.info("Resumen por causal: %s", args.reporte_csv)
    return 0


def _resumir(resumen: Resumen) -> None:
    filas = resumen.filas
    logger.info("─" * 68)
    logger.info("Filas ............ %d en %d factura(s)", len(filas), resumen.facturas)
    logger.info(
        "Llenadas solas ... código %d, clasificación %d, centro de costos %d, gestor %d",
        sum(1 for f in filas if f.codigo_numerico),
        sum(1 for f in filas if f.clasificacion),
        sum(1 for f in filas if f.centro_costos),
        sum(1 for f in filas if f.gestor),
    )
    con = sum(1 for f in filas if f.sugerencia)
    logger.info(
        "Sugerencias ...... %d (%d del criterio del equipo, %d por regla); "
        "%d quedan para el auditor",
        con,
        sum(1 for f in filas if f.confianza == "APRENDIDA"),
        sum(1 for f in filas if f.confianza == "REGLA"),
        len(filas) - con,
    )
    sin_clasif = [f for f in filas if not f.clasificacion]
    if sin_clasif:
        logger.warning(
            "%d fila(s) con una causal que no está en la tabla: %s",
            len(sin_clasif),
            ", ".join(sorted({f.codigo_numerico for f in sin_clasif})[:10]),
        )
    sin_centro = sum(1 for f in filas if not f.centro_costos)
    if sin_centro:
        logger.info("%d fila(s) sin centro de costos propuesto: revisar a mano", sin_centro)
    logger.info("Valor glosado .... $%s", f"{sum(f.valor_glosado for f in filas):,.0f}")


def _escribir_reporte(resumen: Resumen, ruta: Path) -> None:
    por_causal: dict[tuple[str, str], list[FilaMacro]] = defaultdict(list)
    for f in resumen.filas:
        por_causal[(f.codigo_numerico, f.clasificacion)].append(f)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    with ruta.open("w", encoding="utf-8-sig", newline="") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(
            [
                "CAUSAL",
                "CLASIFICACION",
                "FILAS",
                "VALOR_GLOSADO",
                "SUGERENCIA",
                "CONFIANZA",
                "POR_QUE",
                "CAUSAL_TEXTO",
            ]
        )
        for (causal, clasif), fs in sorted(por_causal.items(), key=lambda kv: -len(kv[1])):
            uno = fs[0]
            w.writerow(
                [
                    causal,
                    clasif,
                    len(fs),
                    f"{sum(f.valor_glosado for f in fs):.0f}",
                    uno.sugerencia,
                    uno.confianza,
                    uno.motivo,
                    uno.descripcion_glosa[:120],
                ]
            )


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
