"""respuestas_adres_por_factura.py — el PDF y el Word de respuesta, uno por factura.

De la **macro de respuesta** del paquete saca, para cada factura, los dos
documentos que se radican:

* `RTA_ADRES_<FACTURA>.pdf` — el **REPORTE RTA ADRES**: encabezado con la
  factura, la radicación y el documento del paciente, y la tabla de seis
  columnas con lo que se le respondió a cada glosa.
* `Reporte_Factura_<FACTURA>_<GESTOR>.docx` — el **Word de respuesta**: primero
  lo que se aceptó y después lo demás, una respuesta por párrafo.

El texto de cada respuesta sale **tal cual** de la columna «RTA GLOSA COMPLETA»
de la macro: es lo que redactó el auditor y no se vuelve a armar ni se corrige.

USO (Windows, desde C:\\temp-notas):

    py tools\\respuestas_adres_por_factura.py ^
        --macro  "D:\\...\\RTA GLOSA ADRES PAQ 31068.xlsx" ^
        --salida "D:\\...\\RESPUESTAS_31068"

    REM Solo las facturas de un gestor, cada una en su carpeta:
    py tools\\respuestas_adres_por_factura.py ^
        --macro "D:\\...\\macro.xlsx" --salida "D:\\...\\RESPUESTAS" ^
        --gestor CAROLINA --carpeta-por-factura

Requiere una sola vez: `py -m pip install openpyxl python-docx reportlab`

**Las glosas totales no se responden una por una.** En el reporte del ADRES hay
filas con la «Descripción Glosa» vacía: son el desglose de una reclamación que
el ADRES glosó entera por el FURIPS. No entran en los documentos, pero **sí se
dicen al pie**, con cuántas son y cuánto valen, para no esconder nada. Con
`--incluir-glosas-totales` se meten igual.
"""

from __future__ import annotations

import argparse
import csv
import logging
import sys
from dataclasses import dataclass, field
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from _dinero import a_numero, a_texto  # noqa: E402
from ajustar_detallado_glosas import _abrir_libro, normalizar_factura  # noqa: E402

logger = logging.getLogger("respuestas_adres")

# Columnas de la macro (1 = A). El encabezado va en la fila 2.
COL_RADICACION = 2
COL_FACTURA = 3
COL_PAQUETE = 4
COL_VALOR_GLOSADO = 9
COL_DOC_VICTIMA = 10
COL_DESC_ELEMENTO = 14
COL_DESC_GLOSA = 15
COL_OBSERVACION = 19
COL_RTA_COMPLETA = 21
COL_VALOR_ACEPTADO = 23
COL_GESTOR = 25

ACEPTA = "SE ACEPTA"

# El aviso de extemporaneidad que cierra el Word. NO se pone solo: es una
# afirmación jurídica sobre el paquete y la decide el auditor (`--extemporanea`).
NOTA_EXTEMPORANEA = (
    "ADICIONALMENTE SE INFORMA QUE LA INCONFORMIDAD NO SE RADICO EN LOS TIEMPOS "
    "ESTABLECIDOS SEGÚN RESOLUCION 1236 DE 2023 EN SU ARTICULO 8 NUMERAL 8.5, "
    "POR LO TANTO, CORRESPONDE A UNA GLOSA EXTEMPORANEA."
)

ENCABEZADO_ACEPTADAS = (
    "SE IDENTIFICAN HALLAZGOS Y SE REALIZAN LOS AJUSTES CORRESPONDIENTES "
    "POR GLOSA ACEPTADA PARCIAL POR VALOR DE {valor}"
)


@dataclass
class Glosa:
    """Una fila de la macro, ya leída."""

    valor_glosado: float = 0.0
    descripcion: str = ""
    rta_completa: str = ""
    observacion: str = ""
    valor_aceptado: float = 0.0
    glosa_total: bool = False

    @property
    def aceptada(self) -> bool:
        return self.observacion.upper().startswith(ACEPTA)

    @property
    def sin_decidir(self) -> bool:
        return not self.observacion.strip()


@dataclass
class Factura:
    """Todo lo que la macro dice de una factura."""

    factura: str = ""
    radicacion: str = ""
    documento_paciente: str = ""
    paquete: str = ""
    gestor: str = ""
    glosas: list[Glosa] = field(default_factory=list)

    def responder(self, incluir_totales: bool = False) -> list[Glosa]:
        """Las glosas que van en los documentos, con las aceptadas de primeras."""
        vivas = [g for g in self.glosas if incluir_totales or not g.glosa_total]
        return [g for g in vivas if g.aceptada] + [g for g in vivas if not g.aceptada]

    @property
    def totales_ocultas(self) -> list[Glosa]:
        return [g for g in self.glosas if g.glosa_total]

    @property
    def valor_aceptado(self) -> float:
        return round(sum(g.valor_aceptado for g in self.glosas), 2)

    @property
    def sin_decidir(self) -> int:
        return sum(1 for g in self.glosas if g.sin_decidir and not g.glosa_total)


def leer_macro(ruta: Path, paquete: str | None = None, gestor: str | None = None):
    """Las facturas de la macro, agrupadas y en el orden en que vienen."""
    wb = _abrir_libro(ruta, solo_datos=True, solo_lectura=True)
    try:
        salida: dict[str, Factura] = {}
        for fila in wb.worksheets[0].iter_rows(min_row=2, values_only=True):
            if len(fila) < COL_VALOR_ACEPTADO:
                continue
            numero = str(fila[COL_FACTURA - 1] or "").strip()
            if not numero.upper().startswith("HUS"):
                continue
            paq = str(fila[COL_PAQUETE - 1] or "").strip()
            ges = str(fila[COL_GESTOR - 1] or "").strip()
            if paquete and paq != str(paquete):
                continue
            if gestor and ges.upper() != gestor.strip().upper():
                continue
            fac = salida.setdefault(normalizar_factura(numero), Factura(factura=numero))
            fac.radicacion = fac.radicacion or str(fila[COL_RADICACION - 1] or "").strip()
            fac.documento_paciente = (
                fac.documento_paciente or str(fila[COL_DOC_VICTIMA - 1] or "").strip()
            )
            fac.paquete = fac.paquete or paq
            fac.gestor = fac.gestor or ges
            fac.glosas.append(
                Glosa(
                    valor_glosado=a_numero(fila[COL_VALOR_GLOSADO - 1]),
                    descripcion=str(fila[COL_DESC_ELEMENTO - 1] or "").strip(),
                    rta_completa=str(fila[COL_RTA_COMPLETA - 1] or "").strip(),
                    observacion=str(fila[COL_OBSERVACION - 1] or "").strip(),
                    valor_aceptado=a_numero(fila[COL_VALOR_ACEPTADO - 1]),
                    # Sin causal propia: el ADRES glosó la reclamación entera.
                    glosa_total=not str(fila[COL_DESC_GLOSA - 1] or "").strip(),
                )
            )
        return salida
    finally:
        wb.close()


def _limpio(texto: str, respaldo: str) -> str:
    limpio = "".join(c for c in str(texto or "") if c.isalnum() or c in "-_")
    return limpio or respaldo


def nombre_pdf(factura: str) -> str:
    return f"RTA_ADRES_{_limpio(factura, 'FACTURA')}.pdf"


def nombre_word(factura: str, gestor: str) -> str:
    ges = _limpio(gestor, "")
    return f"Reporte_Factura_{_limpio(factura, 'FACTURA')}{'_' + ges if ges else ''}.docx"


def datos_para_pdf(fac: Factura, incluir_totales: bool = False) -> dict:
    """Traduce la factura al diccionario que espera el servicio del PDF."""
    glosas = fac.responder(incluir_totales)
    ocultas = [] if incluir_totales else fac.totales_ocultas
    return {
        "factura": fac.factura,
        "radicacion": fac.radicacion,
        "documento_paciente": fac.documento_paciente,
        "glosas": [
            {
                "valor_glosado": g.valor_glosado,
                "descripcion": g.descripcion,
                "rta_completa": g.rta_completa,
                "valor_aceptado": g.valor_aceptado,
                "glosa_total": False,
            }
            for g in glosas
        ],
        "resumen": {
            "valor_glosado": round(sum(g.valor_glosado for g in glosas), 2),
            "valor_aceptado": fac.valor_aceptado,
            "glosas_totales_ocultas": len(ocultas),
            "valor_glosas_totales": round(sum(g.valor_glosado for g in ocultas), 2),
            "pendientes": fac.sin_decidir,
        },
    }


def generar_word(
    fac: Factura,
    destino: Path,
    consecutivo: str = "",
    extemporanea: bool = False,
    incluir_totales: bool = False,
) -> None:
    """El Word de respuesta: lo aceptado primero y una respuesta por párrafo."""
    from docx import Document

    doc = Document()
    aceptado = fac.valor_aceptado
    if aceptado > 0:
        cabeza = ENCABEZADO_ACEPTADAS.format(valor=a_texto(aceptado))
        doc.add_paragraph(f"{consecutivo} {cabeza}".strip() if consecutivo else cabeza)

    for g in fac.responder(incluir_totales):
        if g.rta_completa:
            doc.add_paragraph(g.rta_completa)

    ocultas = [] if incluir_totales else fac.totales_ocultas
    if ocultas or fac.sin_decidir or extemporanea:
        doc.add_paragraph("")
    if extemporanea:
        doc.add_paragraph(NOTA_EXTEMPORANEA)
    if ocultas:
        doc.add_paragraph(
            f"NOTA: no se relacionan {len(ocultas)} renglón(es) de GLOSA TOTAL por "
            f"{a_texto(sum(g.valor_glosado for g in ocultas))}: el ADRES glosó la "
            f"reclamación entera por el FURIPS y esos renglones no traen causal "
            f"propia, así que no se responden uno por uno."
        )
    if fac.sin_decidir:
        doc.add_paragraph(
            f"ATENCIÓN: quedan {fac.sin_decidir} glosa(s) sin decidir en esta factura."
        )

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))


@dataclass
class Resultado:
    factura: str
    gestor: str = ""
    pdf: str = ""
    word: str = ""
    glosas: int = 0
    glosas_totales: int = 0
    sin_decidir: int = 0
    valor_glosado: float = 0.0
    valor_aceptado: float = 0.0
    estado: str = "OK"
    observacion: str = ""


def procesar(
    macro: Path,
    salida: Path,
    paquete: str | None = None,
    gestor: str | None = None,
    consecutivo: str = "",
    extemporanea: bool = False,
    incluir_totales: bool = False,
    carpeta_por_factura: bool = False,
) -> list[Resultado]:
    from app.services.evidencia_adres_pdf import generar_pdf_evidencia

    facturas = leer_macro(macro, paquete, gestor)
    if not facturas:
        raise ValueError("La macro no trae ninguna factura (¿es el archivo correcto?).")

    resultados = []
    for n, (_clave, fac) in enumerate(sorted(facturas.items()), 1):
        carpeta = salida / _limpio(fac.factura, "FACTURA") if carpeta_por_factura else salida
        glosas = fac.responder(incluir_totales)
        res = Resultado(
            factura=fac.factura,
            gestor=fac.gestor,
            glosas=len(glosas),
            glosas_totales=len(fac.totales_ocultas),
            sin_decidir=fac.sin_decidir,
            valor_glosado=round(sum(g.valor_glosado for g in glosas), 2),
            valor_aceptado=fac.valor_aceptado,
        )
        try:
            carpeta.mkdir(parents=True, exist_ok=True)
            ruta_pdf = carpeta / nombre_pdf(fac.factura)
            ruta_pdf.write_bytes(generar_pdf_evidencia(datos_para_pdf(fac, incluir_totales)))
            res.pdf = ruta_pdf.name

            ruta_word = carpeta / nombre_word(fac.factura, fac.gestor)
            generar_word(fac, ruta_word, consecutivo, extemporanea, incluir_totales)
            res.word = ruta_word.name
        except Exception as e:  # noqa: BLE001 - una factura mala no tumba el lote
            logger.warning("  %s: %s", fac.factura, e)
            res.estado, res.observacion = "ERROR", str(e)
        resultados.append(res)
        if n % 50 == 0:
            logger.info("  %d/%d…", n, len(facturas))
    return resultados


def escribir_reporte(ruta: Path, resultados: list[Resultado]) -> None:
    ruta.parent.mkdir(parents=True, exist_ok=True)
    with open(ruta, "w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(
            [
                "FACTURA",
                "GESTOR",
                "PDF",
                "WORD",
                "GLOSAS_RESPONDIDAS",
                "GLOSAS_TOTALES_OMITIDAS",
                "SIN_DECIDIR",
                "VALOR_GLOSADO",
                "VALOR_ACEPTADO",
                "ESTADO",
                "OBSERVACION",
            ]
        )
        for r in resultados:
            w.writerow(
                [
                    r.factura,
                    r.gestor,
                    r.pdf,
                    r.word,
                    r.glosas,
                    r.glosas_totales,
                    r.sin_decidir,
                    f"{r.valor_glosado:.2f}",
                    f"{r.valor_aceptado:.2f}",
                    r.estado,
                    r.observacion,
                ]
            )


def construir_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="El PDF y el Word de respuesta, uno por factura.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    p.add_argument("--macro", type=Path, required=True, help="Excel de la macro de respuesta.")
    p.add_argument("--salida", type=Path, required=True, help="Carpeta donde dejarlos.")
    p.add_argument("--paquete", help="Filtrar por número de paquete.")
    p.add_argument("--gestor", help="Solo las facturas de un gestor (CAROLINA, CLAUDIA, OSCAR…).")
    p.add_argument("--consecutivo", default="", help="Consecutivo del oficio, si lleva.")
    p.add_argument(
        "--extemporanea",
        action="store_true",
        help="Cerrar el Word con el aviso de glosa extemporánea (Res. 1236/2023, art. 8, 8.5).",
    )
    p.add_argument(
        "--incluir-glosas-totales",
        action="store_true",
        help="Meter también las filas sin causal propia (por defecto solo se avisan).",
    )
    p.add_argument(
        "--carpeta-por-factura", action="store_true", help="Cada factura en su propia carpeta."
    )
    p.add_argument("--reporte-csv", type=Path, help="Listado de lo generado por factura.")
    return p


def main(argv: list[str] | None = None) -> int:
    args = construir_parser().parse_args(argv)
    logging.basicConfig(level=logging.INFO, format="%(message)s")

    try:
        resultados = procesar(
            args.macro,
            args.salida,
            args.paquete,
            args.gestor,
            args.consecutivo,
            args.extemporanea,
            args.incluir_glosas_totales,
            args.carpeta_por_factura,
        )
    except ValueError as e:
        logger.error("%s", e)
        return 1
    if args.reporte_csv:
        escribir_reporte(args.reporte_csv, resultados)

    ok = [r for r in resultados if r.estado == "OK"]
    errores = [r for r in resultados if r.estado == "ERROR"]
    sin_decidir = [r for r in resultados if r.sin_decidir]

    print(f"\nFacturas          : {len(resultados)}")
    print(f"  documentos       : {len(ok)} PDF y {len(ok)} Word")
    print(f"  glosado          : {a_texto(sum(r.valor_glosado for r in resultados))}")
    print(f"  aceptado         : {a_texto(sum(r.valor_aceptado for r in resultados))}")
    omitidas = sum(r.glosas_totales for r in resultados)
    if omitidas and not args.incluir_glosas_totales:
        print(f"  glosas totales omitidas (se avisan al pie): {omitidas}")
    if sin_decidir:
        print(f"\nOJO — {len(sin_decidir)} factura(s) con glosas SIN DECIDIR en la macro:")
        for r in sorted(sin_decidir, key=lambda x: -x.sin_decidir)[:10]:
            print(f"   {r.factura:<12} {r.sin_decidir} sin decidir")
    if errores:
        print(f"\nCon error ({len(errores)}):")
        for r in errores[:10]:
            print(f"   {r.factura}: {r.observacion}")
    print(f"\nQuedaron en: {args.salida}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
