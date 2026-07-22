"""generar_informe_baja_cartera.py — Informe Word de baja de facturas (Res. 577/2019).

Genera el documento que elabora el ÁREA DE FACTURACIÓN para presentar las
facturas que NO cumplen los parámetros exigidos por la Resolución No. 577 del
16 de octubre de 2019 (manual/reglamento interno de cartera del HUS) para ser
entregadas al área de cartera para su cobro, y que por lo tanto se presentan
para trámite de depuración / baja.

El insumo es el PDF UNIDO que deja el bot UNIR_PDFS.cmd en cada carpeta de
factura (el archivo `_UNIDO_<carpeta>.pdf` o su copia `_UNIDO_<carpeta>.cmd`,
que es el mismo PDF con otra extensión). De ese consolidado se extraen:

- Número de factura, paciente, documento y fechas de atención.
- El VALOR de la factura (para ordenar de la más cara a la más económica).
- El INFORME DE TRABAJO SOCIAL: se localizan sus páginas y se transcribe el
  concepto, y la justificación de cada factura se AJUSTA a lo que trabajo
  social dejó escrito (sin capacidad de pago, no localizable, sin red de
  apoyo, paciente fallecido, extranjero no afiliado, etc.).

El documento resultante:

1. Introducción del área de facturación: deja constancia de que NO fue
   posible remitir estas facturas al área de cartera para su posible cobro
   conforme a la Resolución 577 de 2019, y de que cada una cuenta con
   informe de trabajo social. (NO se incluye la fórmula de "agotadas las
   acciones administrativas de cobro persuasivo (Art. 15)": no hubo cobro
   porque los documentos nunca se remitieron a cartera.)
2. Marco normativo (Resolución 577/2019 — manual de cartera).
3. Relación de facturas ordenadas de MAYOR a MENOR valor.
4. Análisis individual: introducción, lo consignado por trabajo social y
   CONCLUSIÓN de baja para cada factura.
5. Conclusión general y firmas.

USO
---

    # Masivo: una raíz con carpetas de factura (cada una con su _UNIDO_*)
    py generar_informe_baja_cartera.py --raiz "C:\\FACTURAS\\BAJAS" ^
        --salida "C:\\FACTURAS\\BAJAS\\INFORME_BAJA_CARTERA.docx"

    # Ajustes opcionales
    py generar_informe_baja_cartera.py --raiz "C:\\FACTURAS\\BAJAS" ^
        --ciudad Bucaramanga --elaborado "MARIA PEREZ - Facturación"

DEPENDENCIAS
------------
    py -m pip install python-docx pypdf
    (pdfplumber también sirve como lector de PDF si ya está instalado)
"""

from __future__ import annotations

import argparse
import logging
import re
import sys
import unicodedata
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

logger = logging.getLogger("informe_baja_cartera")

RESOLUCION = "Resolución No. 577 del 16 de octubre de 2019"

# ─────────────────────────────────────────────────────────────────────────────
# Utilidades de texto
# ─────────────────────────────────────────────────────────────────────────────


def normalizar(texto: str) -> str:
    if not texto:
        return ""
    t = unicodedata.normalize("NFKD", str(texto))
    t = "".join(c for c in t if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", t).strip().upper()


# Ronda 31: caracteres de control que openpyxl (xlsx) y python-docx (XML)
# rechazan/rompen. Se quitan de todo texto que va a los informes.
_CTRL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def limpiar_parrafo(texto: str) -> str:
    """Colapsa espacios y saltos para citar en el Word sin basura de layout."""
    texto = _CTRL_RE.sub("", texto)
    return re.sub(r"\s+", " ", texto).strip()


def _celda_segura(v):
    """Neutraliza una celda de texto para openpyxl (ronda 31).

    Quita caracteres de control (que tumban el .xlsx) y antepone un
    apóstrofo si el texto empieza con un prefijo de fórmula (=,+,-,@) —
    el texto viene crudo de los PDF y el informe lo abre el auditor en su
    PC, donde Excel ejecutaría =HYPERLINK/=WEBSERVICE/DDE. Los números
    (valores de cartera) pasan intactos.
    """
    if not isinstance(v, str):
        return v
    s = _CTRL_RE.sub("", v)
    if s[:1] in ("=", "+", "-", "@"):
        s = "'" + s
    return s


def extraer_montos(texto: str) -> list[int]:
    montos: list[int] = []
    for m in re.finditer(r"\$?\s*([\d]{1,3}(?:[.,][\d]{3})+(?:[.,]\d{2})?|\d{5,})", texto):
        crudo = m.group(1)
        limpio = crudo
        if re.search(r"[.,]\d{2}$", limpio) and (
            "." in limpio[:-3] or "," in limpio[:-3] or len(limpio) > 6
        ):
            limpio = limpio[:-3]
        limpio = limpio.replace(".", "").replace(",", "")
        if limpio.isdigit() and 3 <= len(limpio) <= 12:
            montos.append(int(limpio))
    return montos


def pesos(valor: int | None) -> str:
    if valor is None:
        return "(valor por confirmar)"
    return f"${valor:,.0f}".replace(",", ".")


# ─────────────────────────────────────────────────────────────────────────────
# Lectura del PDF unido (acepta .pdf y su copia .cmd, que es el mismo PDF)
# ─────────────────────────────────────────────────────────────────────────────


def extraer_paginas_pdf(ruta: Path) -> list[str]:
    """Devuelve el texto por página. Lista vacía si no hay lector o falla."""
    try:
        import pdfplumber

        with pdfplumber.open(ruta) as pdf:
            return [(p.extract_text() or "") for p in pdf.pages]
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"  pdfplumber no pudo leer {ruta.name}: {e}")
        return []
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(ruta))
        return [(p.extract_text() or "") for p in reader.pages]
    except ImportError:
        logger.error("Falta un lector de PDF. Instalar con: py -m pip install pypdf")
        return []
    except Exception as e:
        logger.warning(f"  pypdf no pudo leer {ruta.name}: {e}")
        return []


def hallar_pdf_unido(carpeta: Path) -> Path | None:
    """Busca el consolidado del bot UNIR_PDFS: _UNIDO_*.pdf o _UNIDO_*.cmd."""
    candidatos = sorted(carpeta.glob("_UNIDO_*.pdf")) + sorted(carpeta.glob("_UNIDO_*.cmd"))
    if candidatos:
        return candidatos[0]
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Extracción de datos de la factura desde el texto del PDF unido
# ─────────────────────────────────────────────────────────────────────────────

_RE_FACTURA_TXT = re.compile(
    r"FACTURA\s+ELECTR[OÓ]NICA\s+DE\s+VENTA\s+([A-Z]{2,}\d+)", re.IGNORECASE
)
_RE_FACTURA_NOMBRE = re.compile(r"([A-Z]{2,}0*\d{4,})")
_RE_PACIENTE_FAC = re.compile(r"Paciente\s*(\d{5,16})\s+([A-ZÁÉÍÓÚÑ ]{5,60})", re.IGNORECASE)
_RE_PACIENTE_EPI = re.compile(r"PACIENTE\s*:\s*([A-ZÁÉÍÓÚÑ ]{5,60})", re.IGNORECASE)
_RE_HISTORIA = re.compile(r"HISTORIA\s+CL[IÍ]NICA\s*:?\s*(\d{5,16})", re.IGNORECASE)
_RE_FECHA_INGRESO = re.compile(
    r"FEC(?:HA)?\.?\s*(?:DE\s+)?INGRESO\s*:?\s*([0-9]{1,2}[ /][a-z0-9./ ]{3,12}[0-9]{4})",
    re.IGNORECASE,
)

# Palabras que identifican las páginas del informe de trabajo social
_MARCAS_TSOCIAL = (
    "TRABAJO SOCIAL",
    "TRABAJADORA SOCIAL",
    "TRABAJADOR SOCIAL",
    "VALORACION SOCIAL",
    "INFORME SOCIAL",
    "ESTUDIO SOCIOECONOMICO",
    "VALORACION SOCIOECONOMICA",
)

# Situaciones típicas del concepto de trabajo social → justificación ajustada.
# (patrón regex sobre texto normalizado, etiqueta, párrafo de justificación)
SITUACIONES = [
    (
        r"FALLEC|OBITO|DECESO",
        "paciente fallecido",
        "el paciente falleció, sin que se identificaran deudos o responsables "
        "solidarios con capacidad para asumir la obligación",
    ),
    (
        r"HABITANTE DE CALLE|SITUACION DE CALLE",
        "habitante de calle",
        "el paciente es habitante de calle, sin domicilio, sin red de apoyo "
        "familiar y sin fuente alguna de ingresos verificable",
    ),
    (
        r"NO (?:FUE POSIBLE )?(?:UBICAR|LOCALIZAR|CONTACTAR)|DATOS? (?:ERRAD|FALS|INEXISTENT)|NUMERO(?:S)? (?:ERRAD|EQUIVOCAD)|DIRECCION (?:ERRAD|FALS|INEXISTENT)",
        "paciente/familia no localizable",
        "no fue posible ubicar ni contactar al paciente o a su núcleo familiar "
        "con los datos aportados en la atención, encontrándose datos de "
        "contacto errados o inexistentes",
    ),
    (
        r"SIN (?:CAPACIDAD|RECURSOS)|NO CUENTA CON (?:RECURSOS|INGRESOS|CAPACIDAD)|ESCASOS RECURSOS|BAJOS? RECURSOS|VULNERAB|POBREZA|SISBEN [ABCD]|DESEMPLE",
        "sin capacidad de pago",
        "la valoración socioeconómica evidencia que el paciente y su núcleo "
        "familiar carecen de capacidad de pago para asumir el valor de los "
        "servicios prestados",
    ),
    (
        r"MIGRANTE|EXTRANJER|VENEZOLAN|SIN AFILIACION|NO AFILIAD|NO ASEGURAD",
        "población migrante / sin aseguramiento",
        "se trata de población sin aseguramiento en salud verificable, sin "
        "vínculo con el sistema que permita trasladar el cobro a un tercero "
        "pagador",
    ),
    (
        r"SIN RED DE APOYO|NO (?:TIENE|POSEE|CUENTA CON) (?:RED DE APOYO|FAMILIA)|ABANDONO",
        "sin red de apoyo",
        "el paciente carece de red de apoyo familiar o social que pueda "
        "responder por la obligación",
    ),
    (
        r"SIN (?:DOCUMENTO|IDENTIFICA)|NN|INDOCUMENTAD",
        "paciente sin identificación",
        "el paciente no cuenta con documento de identificación que permita "
        "individualizar un deudor para adelantar gestión de cobro",
    ),
]


def _solo_nombre(crudo: str) -> str:
    """Corta el nombre en el primer token que no venga en MAYÚSCULAS.

    La línea de la factura trae el nombre pegado a la etiqueta siguiente
    ('KEVIN ALEXIS TOBO PEDRAZA TipoExcepcion'): solo se conservan los
    tokens totalmente en mayúsculas.
    """
    tokens = []
    for tok in limpiar_parrafo(crudo).split():
        if tok.isupper():
            tokens.append(tok)
        else:
            break
    return " ".join(tokens)


@dataclass
class Factura:
    carpeta: Path
    numero: str = ""
    paciente: str = ""
    documento: str = ""
    fecha_ingreso: str = ""
    valor: int | None = None
    fuente_pdf: str = ""
    paginas: int = 0
    tsocial_texto: str = ""
    tsocial_paginas: list[int] = field(default_factory=list)
    situaciones: list[tuple[str, str]] = field(default_factory=list)  # (etiqueta, justificación)
    observaciones: list[str] = field(default_factory=list)


_RE_TOTAL_ESTRICTO = re.compile(
    r"TOTAL\s+FACTURA\b|TOTAL\s+NETO|NETO\s+A\s+PAGAR|TOTAL\s+A\s+PAGAR"
    r"|VALOR\s+TOTAL\s+A\s+PAGAR|TOTAL\s+BRUTO\s+FACTURA|VALOR\s+TOTAL\s+ORDEN"
    r"|VALOR\s+NETO\s+FACTURA"
)

# Líneas con "TOTAL" que NO son el total de ESTA factura (acumulados de la
# cuenta del paciente, certificados de facturación, topes SOAT, etc.)
_RE_TOTAL_EXCLUIR = re.compile(
    r"TOTAL\s+FACTURADO|GASTOS\s+A\s*LA\s*FECHA|TOTALDEGASTOS|UVT|SMLV|SMDLV"
)


def _extraer_valor_factura(paginas: list[str]) -> int | None:
    """Valor total de la factura.

    El PDF unido trae de todo (certificados con acumulados de la cuenta,
    facturas de otras IPS, topes SOAT), así que:
    1. Se buscan líneas con etiquetas ESTRICTAS de total de factura y se toma
       el valor MÁS FRECUENTE (la representación DIAN lo repite varias veces).
    2. Si no hay, se aceptan líneas 'TOTAL…' genéricas (excluyendo acumulados).
    3. Último recurso: el mayor monto del documento.
    """
    from collections import Counter

    estrictos: list[int] = []
    genericos: list[int] = []
    for pg in paginas:
        for linea in pg.splitlines():
            ln = normalizar(linea)
            if "TOTAL" not in ln:
                continue
            montos = [m for m in extraer_montos(linea) if m >= 1000]
            if not montos:
                continue
            if _RE_TOTAL_EXCLUIR.search(ln.replace(" ", "")) or _RE_TOTAL_EXCLUIR.search(ln):
                continue
            if _RE_TOTAL_ESTRICTO.search(ln):
                estrictos.extend(montos)
            else:
                genericos.extend(montos)
    for grupo in (estrictos, genericos):
        if grupo:
            return Counter(grupo).most_common(1)[0][0]
    todos = [m for pg in paginas for m in extraer_montos(pg) if m >= 1000]
    return max(todos) if todos else None


def _extraer_trabajo_social(paginas: list[str]) -> tuple[str, list[int]]:
    """Texto de las páginas del informe de trabajo social y sus números (1-based)."""
    nums: list[int] = []
    bloques: list[str] = []
    for i, pg in enumerate(paginas, 1):
        norm = normalizar(pg)
        if any(marca in norm for marca in _MARCAS_TSOCIAL):
            nums.append(i)
            bloques.append(pg)
    if not bloques:
        return "", []
    texto = "\n".join(bloques)
    # Recortar encabezados repetidos de impresión para que la cita sea legible
    lineas = [
        ln
        for ln in texto.splitlines()
        if ln.strip()
        and not re.match(
            r"\s*(Fecha y Hora de Impresi|E\.?S\.?E\.?|Página|Versión|Código:|FOLIO)", ln.strip()
        )
    ]
    return "\n".join(lineas), nums


def analizar_carpeta(carpeta: Path) -> Factura | None:
    f = Factura(carpeta=carpeta)
    pdf = hallar_pdf_unido(carpeta)
    if pdf is None:
        # sin _UNIDO_: usar todos los PDF sueltos de la carpeta como respaldo
        sueltos = sorted(p for p in carpeta.glob("*.pdf"))
        if not sueltos:
            return None
        f.observaciones.append(
            "No se encontró el consolidado _UNIDO_ (bot UNIR_PDFS); se leyeron "
            f"los {len(sueltos)} PDF sueltos de la carpeta."
        )
        paginas = [pg for p in sueltos for pg in extraer_paginas_pdf(p)]
        f.fuente_pdf = f"{len(sueltos)} PDF sueltos"
    else:
        paginas = extraer_paginas_pdf(pdf)
        f.fuente_pdf = pdf.name
    f.paginas = len(paginas)
    if not paginas:
        f.observaciones.append("El PDF no se pudo leer (¿escaneado sin texto o corrupto?).")

    texto_total = "\n".join(paginas)

    # Número de factura: del texto o del nombre de la carpeta
    m = _RE_FACTURA_TXT.search(texto_total)
    if m:
        f.numero = m.group(1).upper()
    else:
        m = _RE_FACTURA_NOMBRE.search(normalizar(carpeta.name).replace(" ", ""))
        f.numero = m.group(1) if m else carpeta.name

    # Paciente y documento
    m = _RE_PACIENTE_FAC.search(texto_total)
    if m:
        f.documento, f.paciente = m.group(1), _solo_nombre(m.group(2))
    else:
        m = _RE_PACIENTE_EPI.search(texto_total)
        if m:
            f.paciente = _solo_nombre(m.group(1))
        m = _RE_HISTORIA.search(texto_total)
        if m:
            f.documento = m.group(1)

    m = _RE_FECHA_INGRESO.search(texto_total)
    if m:
        f.fecha_ingreso = limpiar_parrafo(m.group(1))

    f.valor = _extraer_valor_factura(paginas)
    if f.valor is None:
        f.observaciones.append("No se pudo extraer el valor de la factura del PDF.")

    f.tsocial_texto, f.tsocial_paginas = _extraer_trabajo_social(paginas)
    if f.tsocial_texto:
        norm_ts = normalizar(f.tsocial_texto)
        for patron, etiqueta, justificacion in SITUACIONES:
            if re.search(patron, norm_ts):
                f.situaciones.append((etiqueta, justificacion))
    else:
        f.observaciones.append(
            "NO se encontró el informe de trabajo social dentro del PDF unido: "
            "verificar antes de presentar el documento."
        )
    return f


# ─────────────────────────────────────────────────────────────────────────────
# Generación del Word
# ─────────────────────────────────────────────────────────────────────────────


def generar_word(facturas: list[Factura], salida: Path, ciudad: str, elaborado: str) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        sys.stderr.write("ERROR: falta python-docx. Instalar con: py -m pip install python-docx\n")
        raise SystemExit(1) from None

    hoy = date.today()
    meses = (
        "enero",
        "febrero",
        "marzo",
        "abril",
        "mayo",
        "junio",
        "julio",
        "agosto",
        "septiembre",
        "octubre",
        "noviembre",
        "diciembre",
    )
    fecha_larga = f"{hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"
    total = sum(f.valor or 0 for f in facturas)

    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(11)

    def titulo(texto, nivel=1):
        h = doc.add_heading(texto, level=nivel)
        for run in h.runs:
            run.font.name = "Arial"
        return h

    def parrafo(texto, negrita=False, centrado=False):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = negrita
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrado else WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # ── Encabezado ───────────────────────────────────────────────────────────
    parrafo("E.S.E. HOSPITAL UNIVERSITARIO DE SANTANDER", negrita=True, centrado=True)
    parrafo("ÁREA DE FACTURACIÓN", negrita=True, centrado=True)
    parrafo(
        "INFORME DE FACTURAS QUE NO CUMPLEN LOS PARÁMETROS PARA SU ENTREGA AL "
        "ÁREA DE CARTERA — TRÁMITE DE DEPURACIÓN / BAJA",
        negrita=True,
        centrado=True,
    )
    parrafo(f"{ciudad}, {fecha_larga}", centrado=True)
    doc.add_paragraph()

    # ── 1. Introducción ──────────────────────────────────────────────────────
    titulo("1. INTRODUCCIÓN", 1)
    parrafo(
        "El presente documento es elaborado por el área de facturación de la "
        "E.S.E. Hospital Universitario de Santander con el propósito de "
        "presentar, de manera individual y debidamente soportada, las "
        f"facturas que se relacionan a continuación, las cuales NO cumplen "
        f"con los parámetros y documentos que establece la {RESOLUCION} "
        "(manual para la administración y cobro de la cartera institucional) "
        "para ser entregadas al área de cartera para su cobro."
    )
    parrafo(
        "En consecuencia, el área de facturación deja constancia y justifica "
        "que NO fue posible remitir estas facturas al área de cartera para su "
        f"posible cobro, conforme a lo previsto en la {RESOLUCION}, por cuanto "
        "no se logró conformar el expediente con los documentos que la misma "
        "exige para el efecto. Por tal motivo, sobre estas facturas el área de "
        "cartera no adelantó gestión alguna de cobro."
    )
    parrafo(
        "Cada una de las facturas aquí relacionadas cuenta con el respectivo "
        "INFORME DE TRABAJO SOCIAL, incorporado en el expediente digital "
        "consolidado (PDF unido) de cada cuenta, y el análisis individual de "
        "este documento se ajusta a lo consignado por el profesional de "
        "trabajo social en cada caso."
    )
    parrafo(
        "Con fundamento en lo anterior, se presentan estas facturas para que "
        "se adelante el trámite de depuración y baja a que haya lugar, "
        "ordenadas de la de mayor valor a la de menor valor."
    )

    # ── 2. Marco normativo ───────────────────────────────────────────────────
    titulo("2. MARCO NORMATIVO", 1)
    parrafo(
        f"{RESOLUCION}, por medio de la cual la E.S.E. Hospital Universitario "
        "de Santander adopta el manual para la administración y cobro de la "
        "cartera institucional: establece los parámetros, requisitos y "
        "documentos que deben acompañar las cuentas que el área de "
        "facturación entrega al área de cartera para el inicio de la gestión "
        "de cobro, así como el tratamiento de las cuentas que no reúnen "
        "dichos requisitos."
    )
    parrafo(
        "Las facturas objeto del presente informe se analizan exclusivamente "
        "frente a dichos parámetros documentales y a la situación "
        "socioeconómica de los pacientes consignada en los informes de "
        "trabajo social."
    )

    # ── 3. Relación de facturas (mayor a menor valor) ────────────────────────
    titulo("3. RELACIÓN DE FACTURAS (DE MAYOR A MENOR VALOR)", 1)
    tabla = doc.add_table(rows=1, cols=6)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    encabezados = ("No.", "FACTURA", "PACIENTE", "DOCUMENTO", "VALOR", "SITUACIÓN (TRABAJO SOCIAL)")
    for j, enc in enumerate(encabezados):
        cel = tabla.rows[0].cells[j]
        cel.text = enc
        for p in cel.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, f in enumerate(facturas, 1):
        fila = tabla.add_row().cells
        fila[0].text = str(i)
        fila[1].text = f.numero
        fila[2].text = f.paciente or "(por confirmar)"
        fila[3].text = f.documento or "(por confirmar)"
        fila[4].text = pesos(f.valor)
        fila[5].text = (
            "; ".join(etq for etq, _ in f.situaciones)
            if f.situaciones
            else (
                "ver informe de trabajo social"
                if f.tsocial_texto
                else "SIN INFORME UBICADO — revisar"
            )
        )
    p = doc.add_paragraph()
    r = p.add_run(f"Total facturas: {len(facturas)}    Valor total: {pesos(total)}")
    r.bold = True

    # ── 4. Análisis individual ───────────────────────────────────────────────
    titulo("4. ANÁLISIS INDIVIDUAL DE LAS FACTURAS", 1)
    for i, f in enumerate(facturas, 1):
        titulo(f"4.{i} Factura {f.numero} — {pesos(f.valor)}", 2)

        datos = f"Paciente: {f.paciente or '(por confirmar)'}. Documento: {f.documento or '(por confirmar)'}."
        if f.fecha_ingreso:
            datos += f" Fecha de ingreso/atención: {f.fecha_ingreso}."
        datos += (
            f" Expediente digital: {f.fuente_pdf} ({f.paginas} páginas)." if f.fuente_pdf else ""
        )
        parrafo(datos)

        parrafo(
            f"La factura {f.numero}, por valor de {pesos(f.valor)}, corresponde a "
            "servicios de salud efectivamente prestados por la institución. "
            "Revisado el expediente, la cuenta NO reúne los parámetros y "
            f"documentos exigidos por la {RESOLUCION} para su entrega al área "
            "de cartera, razón por la cual no fue posible remitirla para "
            "gestión de cobro."
        )

        if f.tsocial_texto:
            pags = ", ".join(str(n) for n in f.tsocial_paginas[:6])
            parrafo(
                "Informe de trabajo social (obra en el expediente consolidado, "
                f"página(s) {pags} del PDF unido). El profesional consignó, en lo "
                "pertinente:",
                negrita=True,
            )
            extracto = limpiar_parrafo(f.tsocial_texto)
            if len(extracto) > 1400:
                extracto = extracto[:1400].rsplit(" ", 1)[0] + " […]"
            cita = doc.add_paragraph()
            cita.paragraph_format.left_indent = Pt(24)
            run = cita.add_run(f"“{extracto}”")
            run.italic = True
            if f.situaciones:
                justs = "; ".join(j for _e, j in f.situaciones)
                parrafo(
                    "Ajustado el análisis a lo consignado por trabajo social, se "
                    f"concluye que {justs}, situación que hace inviable la "
                    "gestión de cobro de la presente cuenta."
                )
            else:
                parrafo(
                    "El análisis de la cuenta se ajusta a lo consignado por "
                    "trabajo social en el informe transcrito, del cual se "
                    "desprende la inviabilidad de la gestión de cobro. "
                    "(Revisar y complementar según el caso concreto.)"
                )
        else:
            parrafo(
                "NOTA DE REVISIÓN: en el PDF unido de esta factura no se ubicó "
                "el informe de trabajo social (o el documento está escaneado "
                "sin texto legible). Verificar el expediente y completar este "
                "aparte antes de presentar el informe.",
                negrita=True,
            )

        for obs in f.observaciones:
            parrafo(f"Observación: {obs}")

        parrafo(
            f"CONCLUSIÓN FACTURA {f.numero}: por no cumplir los parámetros "
            f"establecidos en la {RESOLUCION} para su entrega al área de "
            "cartera, y en atención a la situación socioeconómica del "
            "paciente descrita en el informe de trabajo social, se presenta "
            "esta factura para el trámite de depuración y baja de los estados "
            "financieros, previa aprobación de las instancias competentes.",
            negrita=True,
        )
        doc.add_paragraph()

    # ── 5. Conclusión general ────────────────────────────────────────────────
    titulo("5. CONCLUSIÓN GENERAL", 1)
    parrafo(
        f"Se presentan {len(facturas)} facturas por un valor total de "
        f"{pesos(total)}, respecto de las cuales el área de facturación deja "
        "constancia de que no fue posible su remisión al área de cartera para "
        f"gestión de cobro, conforme a la {RESOLUCION}, y que cuentan con los "
        "informes de trabajo social que evidencian la situación de cada "
        "paciente. En consecuencia, se recomienda someter estas cuentas a "
        "consideración de las instancias competentes para su depuración y "
        "baja de los estados financieros de la institución."
    )
    # Ronda 31: el "Valor total" suma f.valor (0 cuando no se pudo extraer del
    # PDF). Dejar constancia de cuántas facturas NO aportan al total para que
    # el documento legal no oculte que hay valores pendientes de confirmar.
    _sin_valor = sum(1 for f in facturas if f.valor is None)
    if _sin_valor:
        parrafo(
            f"NOTA: {_sin_valor} factura(s) se relacionan con el valor "
            "pendiente de confirmar (no fue posible extraerlo del PDF) y NO "
            "están incluidas en el valor total citado; deben verificarse "
            "manualmente antes de la depuración."
        )
    doc.add_paragraph()
    doc.add_paragraph()
    parrafo("_________________________________")
    parrafo(f"Elaboró: {elaborado}")
    parrafo("Área de Facturación — E.S.E. Hospital Universitario de Santander")

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)


# ─────────────────────────────────────────────────────────────────────────────
# Informe Excel (relación de facturas + extractos de trabajo social)
# ─────────────────────────────────────────────────────────────────────────────


def generar_excel(facturas: list[Factura], salida: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        sys.stderr.write("ERROR: falta openpyxl. Instalar con: py -m pip install openpyxl\n")
        raise SystemExit(1) from None

    FILL_HEADER = PatternFill("solid", fgColor="1F4E79")
    FONT_HEADER = Font(color="FFFFFF", bold=True, size=10)
    FILL_OK = PatternFill("solid", fgColor="C6EFCE")
    FONT_OK = Font(color="006100", size=10)
    FILL_MAL = PatternFill("solid", fgColor="FFC7CE")
    FONT_MAL = Font(color="9C0006", bold=True, size=10)
    BORDE = Border(*(Side(style="thin", color="D9D9D9"),) * 4)

    wb = Workbook()

    def hoja(nombre, encabezados, anchos):
        ws = wb.create_sheet(nombre)
        for j, (enc, ancho) in enumerate(zip(encabezados, anchos, strict=False), 1):
            cel = ws.cell(row=1, column=j, value=_celda_segura(enc))
            cel.fill, cel.font = FILL_HEADER, FONT_HEADER
            cel.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            ws.column_dimensions[get_column_letter(j)].width = ancho
        ws.freeze_panes = "A2"
        ws.row_dimensions[1].height = 30
        return ws

    def fila(ws, valores, cols_moneda=(), col_semaforo=None, bueno=True):
        r = ws.max_row + 1
        for j, v in enumerate(valores, 1):
            # v se conserva sin tocar para el chequeo de cols_moneda de abajo;
            # solo la celda escrita se sanea (ronda 31).
            cel = ws.cell(row=r, column=j, value=_celda_segura(v))
            cel.border = BORDE
            cel.alignment = Alignment(vertical="top", wrap_text=True)
            cel.font = Font(size=10)
            if j in cols_moneda and isinstance(v, (int, float)):
                cel.number_format = "#,##0"
                cel.alignment = Alignment(vertical="top", horizontal="right")
        if col_semaforo:
            cel = ws.cell(row=r, column=col_semaforo)
            cel.fill = FILL_OK if bueno else FILL_MAL
            cel.font = FONT_OK if bueno else FONT_MAL

    ws = hoja(
        "RELACION DE FACTURAS",
        [
            "No.",
            "FACTURA",
            "PACIENTE",
            "DOCUMENTO",
            "FECHA ATENCIÓN",
            "VALOR",
            "INFORME TRABAJO SOCIAL",
            "PÁGINAS T. SOCIAL",
            "SITUACIÓN (TRABAJO SOCIAL)",
            "FUENTE (PDF UNIDO)",
            "PÁGINAS PDF",
            "OBSERVACIONES",
            "CONCLUSIÓN",
        ],
        [6, 15, 30, 14, 14, 14, 16, 12, 40, 26, 9, 45, 40],
    )
    wb.active = ws
    del wb["Sheet"]
    for i, f in enumerate(facturas, 1):
        fila(
            ws,
            [
                i,
                f.numero,
                f.paciente or "(por confirmar)",
                f.documento or "(por confirmar)",
                f.fecha_ingreso,
                f.valor,
                "SI" if f.tsocial_texto else "NO — REVISAR",
                ", ".join(str(n) for n in f.tsocial_paginas[:6]),
                "; ".join(etq for etq, _ in f.situaciones),
                f.fuente_pdf,
                f.paginas,
                " | ".join(f.observaciones),
                f"Se presenta para depuración/baja: no cumple los parámetros de la {RESOLUCION} "
                "para entrega al área de cartera.",
            ],
            cols_moneda=(6,),
            col_semaforo=7,
            bueno=bool(f.tsocial_texto),
        )
    total = sum(f.valor or 0 for f in facturas)
    fila(ws, ["", "TOTAL", "", "", "", total, "", "", "", "", "", "", ""], cols_moneda=(6,))
    ws.auto_filter.ref = f"A1:M{ws.max_row - 1}"

    ws = hoja(
        "EXTRACTOS TRABAJO SOCIAL",
        ["FACTURA", "VALOR", "SITUACIÓN(ES)", "EXTRACTO DEL INFORME DE TRABAJO SOCIAL"],
        [15, 14, 35, 120],
    )
    for f in facturas:
        extracto = limpiar_parrafo(f.tsocial_texto)
        if len(extracto) > 2500:
            extracto = extracto[:2500].rsplit(" ", 1)[0] + " […]"
        fila(
            ws,
            [
                f.numero,
                f.valor,
                "; ".join(etq for etq, _ in f.situaciones),
                extracto or "(SIN INFORME DE TRABAJO SOCIAL UBICADO EN EL PDF UNIDO)",
            ],
            cols_moneda=(2,),
        )
    ws.auto_filter.ref = f"A1:D{ws.max_row}"

    ws = hoja("LEYENDA", ["SECCIÓN", "EXPLICACIÓN"], [30, 120])
    for sec, exp in (
        (
            "Qué contiene",
            "Relación de las facturas presentadas para depuración/baja por no cumplir "
            f"los parámetros de la {RESOLUCION} (manual de cartera) para su entrega al "
            "área de cartera, ordenadas de la más cara a la más económica, con la "
            "situación consignada por trabajo social y el extracto de cada informe. "
            "Acompaña al documento Word del mismo nombre.",
        ),
        (
            "INFORME TRABAJO SOCIAL = NO (rojo)",
            "En el PDF unido de esa factura no se ubicó el informe de trabajo social "
            "(o está escaneado sin texto): completar el expediente antes de presentar.",
        ),
        (
            "Herramienta",
            "tools/generar_informe_baja_cartera.py — Motor Glosas HUS. Genera este "
            "Excel y el Word en la misma corrida.",
        ),
    ):
        fila(ws, [sec, exp])

    salida.parent.mkdir(parents=True, exist_ok=True)
    wb.save(salida)


# ─────────────────────────────────────────────────────────────────────────────
# Orquestación
# ─────────────────────────────────────────────────────────────────────────────


def main() -> int:
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument(
        "--raiz",
        type=Path,
        required=True,
        help="Raíz con las carpetas de factura (cada una con su _UNIDO_*.pdf/.cmd)",
    )
    parser.add_argument("--salida", type=Path, default=None, help="Ruta del Word de salida")
    parser.add_argument(
        "--salida-excel",
        type=Path,
        default=None,
        help="Ruta del Excel de salida (por defecto, junto al Word con extensión .xlsx)",
    )
    parser.add_argument("--ciudad", default="Bucaramanga", help="Ciudad del encabezado")
    parser.add_argument(
        "--elaborado",
        default="(nombre y cargo de quien elabora)",
        help="Nombre/cargo para la firma de 'Elaboró'",
    )
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.INFO, format="%(message)s", handlers=[logging.StreamHandler(sys.stdout)]
    )

    if not args.raiz.is_dir():
        logger.error(f"No existe la carpeta: {args.raiz}")
        return 1

    logger.info("=" * 72)
    logger.info(f"INFORME DE BAJA DE FACTURAS — {RESOLUCION} — Motor Glosas HUS")
    logger.info("=" * 72)

    carpetas = sorted(d for d in args.raiz.iterdir() if d.is_dir())
    if hallar_pdf_unido(args.raiz) or any(args.raiz.glob("*.pdf")):
        carpetas.append(args.raiz)

    facturas: list[Factura] = []
    for carpeta in carpetas:
        f = analizar_carpeta(carpeta)
        if f is None:
            continue
        facturas.append(f)
        ts = "con trabajo social" if f.tsocial_texto else "SIN trabajo social"
        logger.info(f"  {f.numero}: {pesos(f.valor)} — {ts} — {f.fuente_pdf}")

    if not facturas:
        logger.error("No se encontró ninguna carpeta con PDF para procesar.")
        return 1

    # De la más cara a la más económica (sin valor van al final)
    facturas.sort(key=lambda f: (f.valor is None, -(f.valor or 0)))

    salida = args.salida or (args.raiz / f"INFORME_BAJA_CARTERA_{date.today():%Y%m%d}.docx")
    generar_word(facturas, salida, args.ciudad, args.elaborado)
    salida_excel = args.salida_excel or salida.with_suffix(".xlsx")
    generar_excel(facturas, salida_excel)

    sin_ts = sum(1 for f in facturas if not f.tsocial_texto)
    logger.info("")
    logger.info("=" * 72)
    logger.info(f"RESULTADO: {len(facturas)} facturas en el informe, ordenadas por valor.")
    if sin_ts:
        logger.info(
            f"⚠ {sin_ts} factura(s) SIN informe de trabajo social ubicado: quedaron "
            "marcadas con NOTA DE REVISIÓN en el Word y en rojo en el Excel."
        )
    logger.info(f"WORD:      {salida}")
    logger.info(f"EXCEL:     {salida_excel}")
    logger.info("=" * 72)
    return 0


if __name__ == "__main__":
    sys.exit(main())
