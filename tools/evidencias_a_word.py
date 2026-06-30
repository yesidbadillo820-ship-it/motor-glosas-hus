"""evidencias_a_word.py — Une los pantallazos de cierre de COOSALUD en un Word.

Lee `.png` de evidencia (típicamente la carpeta del bot de COOSALUD, donde
cada archivo se llama `HUS<numero>_cierre.png`) y arma un .docx con UNA
factura por página: encabezado con el número de factura + la imagen escalada
para que entre completa con el encabezado en una sola hoja A4 + salto de
página.

Acepta dos modos:

  --carpeta  toma todos los archivos que matcheen `--patron` dentro de la
             carpeta (ordenados por nombre).
  --lista    toma rutas explícitas desde un TXT (1 ruta por línea, se
             respeta el orden del TXT). Útil cuando solo querés un subset
             de la carpeta.

USO:
    py evidencias_a_word.py ^
        --carpeta "D:\\USUARIO CARTERA\\Documents\\COOSALUD\\EVIDENCIA" ^
        --salida  "D:\\USUARIO CARTERA\\Documents\\COOSALUD\\evidencias.docx"

    py evidencias_a_word.py ^
        --lista   "D:\\USUARIO CARTERA\\Documents\\COOSALUD\\lote_dia3.txt" ^
        --salida  "D:\\USUARIO CARTERA\\Documents\\COOSALUD\\evidencias_dia3.docx"

DEPENDENCIAS:
    py -m pip install python-docx
"""

from __future__ import annotations

import argparse
import logging
import re
import struct
import sys
from pathlib import Path

logger = logging.getLogger("evidencias_word")

RE_FACTURA = re.compile(r"^(HUS\d{5,9})", re.IGNORECASE)


def _factura_desde_nombre(p: Path) -> str:
    """HUS508259_cierre.png → 'HUS508259'. Si no matchea, usa el stem completo."""
    m = RE_FACTURA.match(p.name)
    return m.group(1).upper() if m else p.stem


def _leer_lista(p: Path) -> list[Path]:
    """Lee un TXT con 1 ruta por linea. Ignora vacias, comentarios (#) y
    quita comillas simples/dobles envolventes. Preserva el orden del TXT."""
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            texto = p.read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise RuntimeError(f"No pude decodificar la lista: {p}")
    rutas: list[Path] = []
    for linea in texto.splitlines():
        s = linea.strip().strip('"').strip("'")
        if not s or s.startswith("#"):
            continue
        rutas.append(Path(s))
    return rutas


def _png_dims_px(path: Path) -> tuple[int, int] | None:
    """Devuelve (ancho_px, alto_px) del PNG leyendo su cabecera IHDR.
    Devuelve None si el archivo no es un PNG válido."""
    try:
        with path.open("rb") as f:
            sig = f.read(8)
            if sig[:4] != b"\x89PNG":
                return None
            f.read(8)  # IHDR length (4) + chunk type (4)
            w, h = struct.unpack(">II", f.read(8))
            return w, h
    except Exception:
        return None


def main() -> int:
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    fuente = parser.add_mutually_exclusive_group(required=True)
    fuente.add_argument("--carpeta", type=Path, help="Carpeta con los PNG de evidencia.")
    fuente.add_argument(
        "--lista", type=Path, help="TXT con 1 ruta de PNG por linea (preserva el orden del TXT)."
    )
    parser.add_argument("--salida", type=Path, required=True, help="Archivo .docx de salida.")
    parser.add_argument(
        "--patron",
        type=str,
        default="*.png",
        help="Glob de archivos (default *.png; ejemplo --patron '*_cierre.png'). Solo aplica con --carpeta.",
    )
    args = parser.parse_args()
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.shared import Cm, Pt
    except ImportError:
        sys.stderr.write("ERROR: falta python-docx. Instalalo con: py -m pip install python-docx\n")
        return 2

    if args.lista is not None:
        if not args.lista.is_file():
            logger.error(f"No es un archivo: {args.lista}")
            return 1
        candidatos = _leer_lista(args.lista)
        if not candidatos:
            logger.error(f"La lista esta vacia: {args.lista}")
            return 1
        imagenes: list[Path] = []
        faltantes: list[Path] = []
        for ruta in candidatos:
            if ruta.is_file():
                imagenes.append(ruta)
            else:
                faltantes.append(ruta)
        if faltantes:
            logger.warning(f"{len(faltantes)} archivos de la lista no existen — se omiten:")
            for f in faltantes:
                logger.warning(f"  - {f}")
        if not imagenes:
            logger.error("Ninguna ruta de la lista existe en disco.")
            return 1
    else:
        if not args.carpeta.is_dir():
            logger.error(f"No es una carpeta: {args.carpeta}")
            return 1
        imagenes = sorted(args.carpeta.glob(args.patron))
        if not imagenes:
            logger.error(f"No encontré imágenes ({args.patron}) en {args.carpeta}")
            return 1
    logger.info(f"Imágenes a incluir: {len(imagenes)}")

    doc = Document()
    # Márgenes razonables para A4.
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)
    # Área útil A4 con esos márgenes: 17.4cm × 26.7cm aprox.
    # Reservamos ~1.8cm para el encabezado (titulo 20pt + spacing): la imagen
    # no debe pasarse de ANCHO_MAX × ALTO_MAX para que entre JUNTO al encabezado
    # en la misma página, sin que Word la empuje a la siguiente hoja.
    ANCHO_MAX_CM = 17.0
    ALTO_MAX_CM = 22.0

    for i, img in enumerate(imagenes):
        factura = _factura_desde_nombre(img)
        # Encabezado: número de factura, grande, negrita, centrado.
        # Quitamos el space_after para que la imagen quede pegada y no fuerce
        # un overflow a la página siguiente.
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(factura)
        run.bold = True
        run.font.size = Pt(20)
        # Decidir width/height para que la imagen ENTRE en una sola hoja
        # JUNTO con el encabezado (sin salto de página automático).
        dims = _png_dims_px(img)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        try:
            if dims is None:
                # Fallback: que Word elija — limitamos solo el ancho.
                p_img.add_run().add_picture(str(img), width=Cm(ANCHO_MAX_CM))
            else:
                w_px, h_px = dims
                if w_px <= 0 or h_px <= 0:
                    p_img.add_run().add_picture(str(img), width=Cm(ANCHO_MAX_CM))
                else:
                    # Si fijar ancho=ANCHO_MAX_CM excede ALTO_MAX_CM,
                    # mejor fijamos por alto (la imagen queda más angosta
                    # pero entra completa con el encabezado).
                    alto_si_max_ancho = ANCHO_MAX_CM * (h_px / w_px)
                    if alto_si_max_ancho <= ALTO_MAX_CM:
                        p_img.add_run().add_picture(str(img), width=Cm(ANCHO_MAX_CM))
                    else:
                        p_img.add_run().add_picture(str(img), height=Cm(ALTO_MAX_CM))
        except Exception as e:
            logger.warning(f"  {img.name}: no se pudo insertar — {e}")
            continue
        # Salto de página (menos en la última).
        if i < len(imagenes) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        logger.info(f"  [{i + 1}/{len(imagenes)}] {factura} ← {img.name}")

    args.salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.salida))
    logger.info(f"\nWord generado: {args.salida}")
    logger.info(f"  Páginas: {len(imagenes)} (una factura por página)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
